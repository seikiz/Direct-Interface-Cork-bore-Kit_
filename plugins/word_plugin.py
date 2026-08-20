# word_plugin.py - Word文档导出插件（集成式）
import os
import re
from datetime import datetime
from plugin_base import PluginBase

_docx_loaded = False


def _ensure_docx():
    """惰性导入：首次生成 Word 时才加载 python-docx"""
    global _docx_loaded, Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT
    if not _docx_loaded:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        _docx_loaded = True


def _docx_available():
    try:
        import importlib.util
        return importlib.util.find_spec("docx") is not None
    except Exception:
        return False


class WordPlugin(PluginBase):
    name = "Word导出"
    version = "1.0"
    description = "生成Word文档：公文、报告、信函等"
    author = "seiki"
    enabled = True

    def __init__(self, core):
        super().__init__(core)
        self.export_dir = "exports"
        os.makedirs(self.export_dir, exist_ok=True)

    def on_command(self, command, args):
        if command == "doc":
            if not _docx_available():
                return "⚠️ 请安装 python-docx: pip install python-docx", False
            return self._handle(args), False
        elif command == "doc_help":
            return self._help(), False
        return None

    def _help(self):
        lines = [
            "📋 Word导出插件命令：",
            "/doc <模板> [参数]",
            "",
            "【模板】",
            "  letter      - 信函/书信",
            "  report      - 工作报告",
            "  meeting     - 会议纪要",
            "  notice      - 通知/公告",
            "  approval    - 审批单",
            "  summary     - 总结报告",
            "  blank       - 空白文档",
            "",
            "示例:",
            "/doc letter 收件人=张主任 内容=会议通知",
            "/doc report 标题=项目进展 内容=已完成阶段一",
            "/doc meeting 议题=项目会 决议=通过方案A",
            "/doc notice 标题=放假通知 内容=8月15日放假一天",
        ]
        return "\n".join(lines)

    def _handle(self, args):
        _ensure_docx()
        if not args:
            return self._help()

        parts = args.strip().split(maxsplit=1)
        template = parts[0].lower()
        data = self._parse(parts[1] if len(parts) > 1 else "")

        templates = {
            "letter": self._gen_letter,
            "report": self._gen_report,
            "meeting": self._gen_meeting,
            "notice": self._gen_notice,
            "approval": self._gen_approval,
            "summary": self._gen_summary,
            "blank": self._gen_blank,
        }

        func = templates.get(template)
        if not func:
            return f"未知模板: {template}\n{self._help()}"

        try:
            return func(data)
        except Exception as e:
            return f"生成失败: {e}"

    def _parse(self, raw):
        result = {}
        for m in re.finditer(r'(\w+)="(.*?)"', raw):
            result[m.group(1)] = m.group(2)
        for m in re.finditer(r'(\w+)=(\S+)', raw):
            if m.group(1) not in result:
                result[m.group(1)] = m.group(2)
        return result

    def _save(self, doc, name):
        try:
            import doc_layout
            doc_layout.apply_official_format(doc)  # 公文排版算法
        except Exception as e:
            print(f"[Word导出] 排版失败: {e}")
        fn = f"{name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path = os.path.join(self.export_dir, fn)
        doc.save(path)
        return f"✅ {name}已生成: {path}"

    def _add_heading(self, doc, text, level=1):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def _add_paragraph(self, doc, text, bold=False, align=None):
        p = doc.add_paragraph(text)
        if bold:
            p.runs[0].bold = True if p.runs else False
        if align:
            p.alignment = align
        return p

    # ========== 模板 ==========

    def _gen_letter(self, d):
        doc = Document()
        # 标题
        self._add_heading(doc, d.get("标题", "信函"), level=1)
        # 收件人
        doc.add_paragraph(f"{d.get('收件人', '')}：")
        # 内容
        doc.add_paragraph(d.get("内容", ""))
        # 落款
        doc.add_paragraph()
        doc.add_paragraph(f"{d.get('落款', '')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        doc.add_paragraph(d.get("日期", datetime.now().strftime("%Y年%m月%d日")), align=WD_ALIGN_PARAGRAPH.RIGHT)
        return self._save(doc, f"信函_{d.get('收件人','')}")

    def _gen_report(self, d):
        doc = Document()
        self._add_heading(doc, d.get("标题", "工作报告"), level=1)
        doc.add_paragraph(f"报告人：{d.get('报告人', '')}")
        doc.add_paragraph(f"日期：{d.get('日期', datetime.now().strftime('%Y年%m月%d日'))}")
        doc.add_paragraph()
        doc.add_paragraph(d.get("内容", ""))
        return self._save(doc, "工作报告")

    def _gen_meeting(self, d):
        doc = Document()
        self._add_heading(doc, d.get("标题", "会议纪要"), level=1)
        doc.add_paragraph(f"时间：{d.get('时间', '')}")
        doc.add_paragraph(f"地点：{d.get('地点', '')}")
        doc.add_paragraph(f"主持：{d.get('主持', '')}")
        doc.add_paragraph(f"记录：{d.get('记录', '')}")
        doc.add_paragraph(f"参会：{d.get('参会', '')}")
        doc.add_paragraph()
        doc.add_paragraph("议题：", style='List Bullet')
        for item in d.get("议题", "").split('\n'):
            if item.strip():
                doc.add_paragraph(item.strip(), style='List Number')
        doc.add_paragraph()
        doc.add_paragraph("决议：")
        for item in d.get("决议", "").split('\n'):
            if item.strip():
                doc.add_paragraph(f"• {item.strip()}")
        return self._save(doc, "会议纪要")

    def _gen_notice(self, d):
        doc = Document()
        self._add_heading(doc, d.get("标题", "通知"), level=1)
        doc.add_paragraph(d.get("内容", ""))
        doc.add_paragraph()
        doc.add_paragraph(f"{d.get('发布单位', '')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        doc.add_paragraph(d.get("日期", datetime.now().strftime("%Y年%m月%d日")), align=WD_ALIGN_PARAGRAPH.RIGHT)
        return self._save(doc, "通知")

    def _gen_approval(self, d):
        doc = Document()
        self._add_heading(doc, "审批单", level=1)
        doc.add_paragraph(f"申请人：{d.get('申请人', '')}")
        doc.add_paragraph(f"申请部门：{d.get('部门', '')}")
        doc.add_paragraph(f"申请日期：{d.get('日期', datetime.now().strftime('%Y年%m月%d日'))}")
        doc.add_paragraph()
        doc.add_paragraph(f"申请事由：{d.get('事由', '')}")
        doc.add_paragraph()
        doc.add_paragraph("审批意见：")
        doc.add_paragraph()
        doc.add_paragraph("审批人签字：__________________")
        doc.add_paragraph("日期：__________________")
        return self._save(doc, f"审批单_{d.get('申请人','')}")

    def _gen_summary(self, d):
        doc = Document()
        self._add_heading(doc, d.get("标题", "工作总结"), level=1)
        doc.add_paragraph(f"部门：{d.get('部门', '')}")
        doc.add_paragraph(f"时间：{d.get('时间', '')}")
        doc.add_paragraph()
        doc.add_paragraph("一、基本情况")
        doc.add_paragraph(d.get("基本情况", ""))
        doc.add_paragraph()
        doc.add_paragraph("二、主要工作")
        doc.add_paragraph(d.get("主要工作", ""))
        doc.add_paragraph()
        doc.add_paragraph("三、存在问题")
        doc.add_paragraph(d.get("存在问题", ""))
        doc.add_paragraph()
        doc.add_paragraph("四、下一步计划")
        doc.add_paragraph(d.get("下一步", ""))
        return self._save(doc, "工作总结")

    def _gen_blank(self, d):
        doc = Document()
        self._add_heading(doc, d.get("标题", "空白文档"), level=1)
        for i in range(int(d.get("段落数", "3"))):
            doc.add_paragraph(f"（第{i+1}段）")
        return self._save(doc, "空白文档")