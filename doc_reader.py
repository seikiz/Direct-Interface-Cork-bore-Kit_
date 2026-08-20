# ============================================================
#   doc_reader.py - 文档读入（Word/Excel → 纯文本）
#
#   支持 .docx（段落+表格）与 .xlsx（多工作表逐行），
#   提取结果交给 AI 做总结/改写/翻译/问答。
# ============================================================

import os


def read_docx(path):
    """读取 Word：段落 + 表格（每表以【表格N】标注）"""
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for ti, table in enumerate(doc.tables, 1):
        parts.append(f"【表格{ti}】")
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_xlsx(path):
    """读取 Excel：逐工作表逐行，以「工作表：名称」分节"""
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"【工作表：{ws.title}】")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v).strip() for v in row]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_document(path):
    """按扩展名读取文档为纯文本"""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx(path)
    if ext == ".xlsx":
        return read_xlsx(path)
    raise ValueError(f"暂不支持的文件类型：{ext}（支持 .docx / .xlsx）")
