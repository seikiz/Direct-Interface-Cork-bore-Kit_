# -*- coding: utf-8 -*-
# ============================================================
#   html_app.py - HTML 前端（pywebview 壳）
#   后端复用 ChatCore + 插件体系；前端 web/index.html（纯 HTML/CSS/JS）
#   桥接：JS 轮询 api.poll() 拉增量消息 / 流式文本，调用 api.* 触发动作
#
#   架构标识：Direct-Interface Cork-bore Kit (DICK) — CODEX engine
#   本文件及本项目的核心架构（树状记忆 / 机制卡 / 战斗 / CODEX 打包）
#   均受 dick_mark.py 溯源水印保护；抄袭者无法剥离架构签名。
# ============================================================

import base64
import json
import os
import sys
import threading
import re

# PyInstaller windowed 模式（无控制台）下 stdout/stderr 为 None，print 会崩溃。
# 启动即重定向到 exe 旁的 debug.log。
if sys.stdout is None or sys.stderr is None:
    _exe = sys.executable if getattr(sys, "frozen", False) else os.path.abspath(__file__)
    _log_path = os.path.join(os.path.dirname(_exe), "debug.log")
    try:
        _log = open(_log_path, "w", encoding="utf-8", errors="replace")
        if sys.stdout is None:
            sys.stdout = _log
        if sys.stderr is None:
            sys.stderr = _log
    except Exception:
        pass

import app_paths
import i18n
import card_compat
import save_guard
import codex_core
from DICK_core import ChatCore
from plugin_manager import PluginManager

BASE_DIR = app_paths.get_base_dir()

# 内置代理通道（中转）：固定地址 → Worker → 隧道 → 本地 net.py → 真实厂商
BUILTIN_RELAY = "https://dick-workshop.seiki342008.workers.dev"

# ---------- 角色卡结构化字段 / 世界卡参数（精细化创作） ----------
ROLE_FIELDS = [
    ("appearance", "外貌"),
    ("personality", "性格"),
    ("background", "过去经历"),
    ("speech", "说话方式"),
    ("first_mes", "开场白"),
    ("mes_example", "对话示例"),
    ("notes", "备注"),
]

WORLD_PARAMS = [
    ("tech_level", "科技水平"),
    ("supernatural", "超自然体系"),
    ("physics", "物理法则"),
    ("time_flow", "时间流速"),
    ("climate", "气候环境"),
    ("geography", "地理格局"),
    ("politics", "政治格局"),
    ("economy", "经济体系"),
]


def html_app_clean_battle(battle):
    """规整战斗配置（类型安全）：attrs / mech_attrs / formulas / moves / buffs"""
    if not isinstance(battle, dict):
        return None
    out = {}
    attrs = battle.get("attrs")
    if isinstance(attrs, dict):
        cleaned = {}
        for key, a in attrs.items():
            if not isinstance(a, dict):
                continue
            key = str(key).strip()[:20]
            if not key:
                continue
            entry = {"label": str(a.get("label") or key).strip()[:20]}
            try:
                entry["initial"] = int(a.get("initial", 10) or 10)
            except (TypeError, ValueError):
                entry["initial"] = 10
            if key == "hp":
                try:
                    entry["max"] = int(a.get("max", 100) or 100)
                except (TypeError, ValueError):
                    entry["max"] = 100
                entry["initial"] = min(entry.get("initial", 100), entry["max"])
            cleaned[key] = entry
        if cleaned:
            out["attrs"] = cleaned
    mech_attrs = battle.get("mech_attrs")
    if isinstance(mech_attrs, list):
        cleaned = []
        for a in mech_attrs:
            if not isinstance(a, dict) or not str(a.get("key") or "").strip():
                continue
            key = str(a["key"]).strip()[:20]
            entry = {"key": key, "label": str(a.get("label") or key).strip()[:20]}
            try:
                entry["initial"] = int(a.get("initial", 10) or 10)
            except (TypeError, ValueError):
                entry["initial"] = 10
            try:
                entry["max"] = int(a.get("max", 999999) or 999999)
            except (TypeError, ValueError):
                entry["max"] = 999999
            cleaned.append(entry)
        if cleaned:
            out["mech_attrs"] = cleaned
    formulas = battle.get("formulas")
    if isinstance(formulas, dict):
        cleaned = {}
        for k, v in formulas.items():
            k = str(k).strip()[:30]
            if k and isinstance(v, str) and v.strip():
                cleaned[k] = v.strip()[:200]
        if cleaned:
            out["formulas"] = cleaned
    moves = battle.get("moves")
    if isinstance(moves, list):
        cleaned = []
        for m in moves:
            if not isinstance(m, dict) or not str(m.get("id") or "").strip():
                continue
            entry = {"id": str(m["id"]).strip()[:30],
                     "name": str(m.get("name") or m["id"]).strip()[:30],
                     "desc": str(m.get("desc") or "").strip()[:60]}
            if isinstance(m.get("formula"), str) and m["formula"].strip():
                entry["formula"] = m["formula"].strip()[:200]
            if isinstance(m.get("cost"), dict):
                cost = {}
                for k, v in m["cost"].items():
                    try:
                        cost[str(k).strip()[:20]] = int(v)
                    except (TypeError, ValueError):
                        pass
                if cost:
                    entry["cost"] = cost
            if isinstance(m.get("buffs"), list):
                bl = []
                for b in m["buffs"]:
                    if isinstance(b, dict) and b.get("id"):
                        try:
                            turns = int(b.get("turns", 3) or 3)
                        except (TypeError, ValueError):
                            turns = 3
                        bl.append({"id": str(b["id"]).strip()[:30], "turns": max(1, turns)})
                if bl:
                    entry["buffs"] = bl
            cleaned.append(entry)
        if cleaned:
            out["moves"] = cleaned
    buffs = battle.get("buffs")
    if isinstance(buffs, list):
        cleaned = []
        for b in buffs:
            if not isinstance(b, dict) or not str(b.get("id") or "").strip():
                continue
            entry = {"id": str(b["id"]).strip()[:30],
                     "name": str(b.get("name") or b["id"]).strip()[:30],
                     "desc": str(b.get("desc") or "").strip()[:60]}
            try:
                entry["turns"] = int(b.get("turns", 3) or 3)
            except (TypeError, ValueError):
                entry["turns"] = 3
            if isinstance(b.get("attrs"), dict):
                at = {}
                for k, v in b["attrs"].items():
                    try:
                        at[str(k).strip()[:20]] = int(v)
                    except (TypeError, ValueError):
                        pass
                if at:
                    entry["attrs"] = at
            cleaned.append(entry)
        if cleaned:
            out["buffs"] = cleaned
    return out or None


def _parse_fields(fields_json):
    """fields_json 可能是 dict、JSON 字符串或旧版纯文本设定；返回 dict"""
    if fields_json is None:
        return {}
    if isinstance(fields_json, dict):
        return fields_json
    s = str(fields_json)
    try:
        d = json.loads(s)
        if isinstance(d, dict):
            return d
    except Exception:
        pass
    return {}


def html_app_clean_mechanics(mech):
    """规整机制卡字段（类型安全）：affection / status.fields / events"""
    if not isinstance(mech, dict):
        return None
    out = {}
    aff = mech.get("affection")
    if isinstance(aff, dict) and aff.get("enabled"):
        try:
            mn = int(aff.get("min", 0) or 0)
        except (TypeError, ValueError):
            mn = 0
        try:
            mx = int(aff.get("max", 100) or 100)
        except (TypeError, ValueError):
            mx = 100
        if mn >= mx:
            mn, mx = 0, 100
        try:
            initial = int(aff.get("initial", 50) or 50)
        except (TypeError, ValueError):
            initial = 50
        initial = max(mn, min(mx, initial))
        try:
            crit = float(aff.get("crit", 0.001) or 0.001)
        except (TypeError, ValueError):
            crit = 0.001
        out["affection"] = {"enabled": True, "initial": initial, "min": mn, "max": mx,
                            "crit": max(0.0, min(1.0, crit))}
    status = mech.get("status")
    if isinstance(status, dict) and status.get("enabled"):
        fields = []
        for f in (status.get("fields") or []):
            if not isinstance(f, dict) or not str(f.get("key") or "").strip():
                continue
            key = str(f["key"]).strip()[:20]
            ftype = "int" if f.get("type") == "int" else "enum"
            field = {"key": key, "name": str(f.get("name") or key).strip()[:20], "type": ftype}
            if ftype == "int":
                try:
                    mn = int(f.get("min", 0) or 0)
                except (TypeError, ValueError):
                    mn = 0
                try:
                    mx = int(f.get("max", 100) or 100)
                except (TypeError, ValueError):
                    mx = 100
                if mn >= mx:
                    mn, mx = 0, 100
                try:
                    initial = int(f.get("initial", 0) or 0)
                except (TypeError, ValueError):
                    initial = 0
                field.update({"min": mn, "max": mx, "initial": max(mn, min(mx, initial))})
            else:
                opts = f.get("options")
                if isinstance(opts, str):
                    opts = [o.strip() for o in opts.replace("，", ",").split(",") if o.strip()]
                if isinstance(opts, list):
                    opts = [str(o).strip()[:20] for o in opts if str(o).strip()]
                field["options"] = opts if opts else []
                field["initial"] = str(f.get("initial") or (opts[0] if opts else "")).strip()[:20]
            fields.append(field)
        if fields:
            out["status"] = {"enabled": True, "fields": fields}
    events = mech.get("events")
    if isinstance(events, list):
        evs = []
        for ev in events:
            if not isinstance(ev, dict) or not str(ev.get("id") or "").strip():
                continue
            e = {"id": str(ev["id"]).strip()[:30],
                 "name": str(ev.get("name") or ev["id"]).strip()[:30],
                 "prompt": str(ev.get("prompt") or "").strip()}
            try:
                if ev.get("aff_ge") is not None:
                    e["aff_ge"] = int(ev["aff_ge"])
            except (TypeError, ValueError):
                pass
            try:
                if ev.get("aff_le") is not None:
                    e["aff_le"] = int(ev["aff_le"])
            except (TypeError, ValueError):
                pass
            kws = ev.get("keywords")
            if isinstance(kws, str):
                kws = [x.strip() for x in kws.replace("，", ",").split(",") if x.strip()]
            if isinstance(kws, list):
                kws = [str(k).strip()[:20] for k in kws if str(k).strip()]
            if kws:
                e["keywords"] = kws
            if ev.get("once", True):
                e["once"] = True
            evs.append(e)
        if evs:
            out["events"] = evs
    roll = mech.get("roll")
    if isinstance(roll, dict):
        cleaned_roll = {}
        _DEF = {"crit": 10, "rare": 4, "fail": 2, "chosen": 0.1, "collapse": 0.001}
        for k in _DEF:
            try:
                v = float(roll.get(k, _DEF[k]))
            except (TypeError, ValueError):
                v = _DEF[k]
            cleaned_roll[k] = max(0.0, min(100.0, v))
        out["roll"] = cleaned_roll
    return out or None


def assemble_role_prompt(name, fields, legacy):
    """由结构化字段拼装系统提示：旧版完整设定保留在前，分节字段追加在后"""
    parts = []
    legacy = (legacy or "").strip()
    if legacy:
        parts.append(legacy)
    sections = []
    for key, label in ROLE_FIELDS:
        v = fields.get(key) or ""
        if isinstance(v, list):
            v = "、".join(str(x) for x in v if str(x).strip())
        v = str(v).strip()
        if v:
            sections.append("【" + label + "】" + chr(10) + v)
    if sections:
        if parts:
            parts.append(chr(10).join(sections))
        else:
            parts.append("你现在的身份是：" + (name or "未命名") + "。" + chr(10) + chr(10) + chr(10).join(sections))
    return chr(10) + chr(10).join(parts) if len(parts) > 1 else (parts[0] if parts else "")


def _render_world_desc(w):
    """把世界卡渲染给 core：背景 + 世界参数（人类可读标签）"""
    desc = str(w.get("description") or "").strip()
    params = w.get("params") or {}
    plines = []
    for key, label in WORLD_PARAMS:
        v = str(params.get(key) or "").strip()
        if v:
            plines.append(label + "：" + v)
    if plines:
        if desc:
            desc += chr(10) + "【世界参数】" + "；".join(plines)
        else:
            desc = "【世界参数】" + "；".join(plines)
    return desc


# ---------- 模型商目录（14 家 / 100+ 模型：选厂商→选模型→跳官网） ----------
PROVIDERS = [
    {"id": "deepseek", "name": "DeepSeek 官方", "free": False,
     "base_url": "https://api.deepseek.com",
     "models": ["deepseek-v4-flash", "deepseek-v4-pro", "deepseek-chat", "deepseek-reasoner"],
     "buy_url": "https://platform.deepseek.com/"},
    {"id": "ovh", "name": "OVH 免费链（免 Key）", "free": True,
     "base_url": "https://oai.endpoints.kepler.ai.cloud.ovh.net/v1",
     "models": ["Qwen3.5-397B-A17B", "Qwen3.6-27B", "Qwen2.5-VL-72B-Instruct",
                "Mistral-Small-3.2-24B-Instruct-2506", "Llama-3.3-70B-Instruct",
                "DeepSeek-R1-Distill-Llama-70B", "Qwen3.5-9B", "Mistral-7B-Instruct-v0.3"],
     "buy_url": "https://endpoints.ai.cloud.ovh.net/"},
    {"id": "alibaba", "name": "阿里云百炼（通义千问）", "free": False,
     "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
     "models": ["qwen-max", "qwen-plus", "qwen-turbo", "qwen-long", "qwen-flash",
                "qwen3-235b-a22b", "qwen3-32b", "qwen3-30b-a3b", "qwen3-14b", "qwen3-8b",
                "qwen2.5-72b-instruct", "qwen2.5-coder-32b-instruct",
                "qwen-vl-max", "qwen-vl-plus"],
     "buy_url": "https://bailian.console.aliyun.com/"},
    {"id": "zhipu", "name": "智谱 AI（GLM）", "free": False,
     "base_url": "https://open.bigmodel.cn/api/paas/v4",
     "models": ["glm-4.6", "glm-4.5-air", "glm-4-plus", "glm-4-air",
                "glm-4-flash", "glm-4-long", "glm-4v-plus", "glm-4.5v"],
     "buy_url": "https://open.bigmodel.cn/"},
    {"id": "siliconflow", "name": "硅基流动 SiliconFlow", "free": False,
     "base_url": "https://api.siliconflow.cn/v1",
     "models": ["deepseek-ai/DeepSeek-V3", "deepseek-ai/DeepSeek-V3.2-Exp", "deepseek-ai/DeepSeek-R1",
                "Qwen/Qwen3-235B-A22B", "Qwen/Qwen3-32B", "Qwen/Qwen3-30B-A3B", "Qwen/Qwen3-14B",
                "Qwen/Qwen2.5-72B-Instruct", "Qwen/Qwen2.5-Coder-32B-Instruct",
                "Qwen/Qwen2.5-VL-72B-Instruct", "zai-org/GLM-4.5-Air", "moonshotai/Kimi-K2-Instruct"],
     "buy_url": "https://siliconflow.cn/"},
    {"id": "moonshot", "name": "Moonshot Kimi", "free": False,
     "base_url": "https://api.moonshot.cn/v1",
     "models": ["kimi-latest", "kimi-k2-0711-preview", "kimi-k2-turbo-preview", "kimi-thinking-preview",
                "moonshot-v1-8k", "moonshot-v1-32k", "moonshot-v1-128k"],
     "buy_url": "https://platform.moonshot.cn/"},
    {"id": "volcengine", "name": "火山方舟（豆包）", "free": False,
     "base_url": "https://ark.cn-beijing.volces.com/api/v3",
     "models": ["doubao-1-5-pro-32k-250115", "doubao-1-5-lite-32k-250115",
                "doubao-pro-32k", "doubao-lite-32k", "doubao-pro-256k",
                "deepseek-v3-241226", "deepseek-r1-250120"],
     "buy_url": "https://console.volcengine.com/ark"},
    {"id": "baidu", "name": "百度千帆（文心）", "free": False,
     "base_url": "https://qianfan.baidubce.com/v2",
     "models": ["ernie-4.0-turbo-8k", "ernie-4.0-8k", "ernie-4.5-8k-preview",
                "ernie-3.5-8k", "ernie-speed-8k", "ernie-lite-8k"],
     "buy_url": "https://console.bce.baidu.com/qianfan"},
    {"id": "minimax", "name": "MiniMax", "free": False,
     "base_url": "https://api.minimax.chat/v1",
     "models": ["MiniMax-Text-01", "abab6.5s-chat", "abab6.5g-chat"],
     "buy_url": "https://platform.minimaxi.com"},
    {"id": "stepfun", "name": "阶跃星辰 StepFun", "free": False,
     "base_url": "https://api.stepfun.com/v1",
     "models": ["step-2-16k", "step-1-8k", "step-1-32k", "step-1-128k", "step-1v-8k"],
     "buy_url": "https://platform.stepfun.com"},
    {"id": "openai", "name": "OpenAI", "free": False,
     "base_url": "https://api.openai.com/v1",
     "models": ["gpt-4o", "gpt-4o-mini", "gpt-4.1", "gpt-4.1-mini", "gpt-4.1-nano",
                "o3", "o3-mini", "o4-mini", "chatgpt-4o-latest"],
     "buy_url": "https://platform.openai.com/"},
    {"id": "anthropic", "name": "Anthropic Claude", "free": False,
     "base_url": "https://api.anthropic.com/v1",
     "models": ["claude-opus-4-20250514", "claude-sonnet-4-20250514",
                "claude-3-7-sonnet-latest", "claude-3-5-sonnet-latest",
                "claude-3-5-haiku-latest", "claude-3-opus-latest"],
     "buy_url": "https://console.anthropic.com/"},
    {"id": "gemini", "name": "Google Gemini", "free": False,
     "base_url": "https://generativelanguage.googleapis.com/v1beta/openai",
     "models": ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.5-flash-lite",
                "gemini-2.0-flash", "gemini-2.0-flash-lite", "gemini-2.0-flash-thinking-exp",
                "gemini-1.5-pro", "gemini-1.5-flash"],
     "buy_url": "https://aistudio.google.com/"},
    {"id": "ollama", "name": "Ollama 本地（免费）", "free": True,
     "base_url": "http://localhost:11434/v1",
     "models": ["qwen3:32b", "qwen3:14b", "qwen3:8b", "qwen2.5:14b",
                "llama3.3:70b", "llama3.1:8b", "deepseek-r1:32b", "deepseek-r1:14b",
                "glm4:9b", "phi4:14b", "gemma3:12b", "mistral:7b"],
     "buy_url": "https://ollama.com/"},
]

# 指令模板：各模型家族默认停止序列（设置里可覆盖；预设 stop_sequences 次之）
_MODEL_STOP_DEFAULTS = {
    "ollama": ["<|im_end|>", "</s>"],
    "qwen": ["<|im_end|>"],
}
for _p in PROVIDERS:
    _p.setdefault("stop", _MODEL_STOP_DEFAULTS.get(_p["id"], []))


def _web_root():
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, "web")
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), "web")


def _seed_defaults(base_dir):
    """exe 首次运行：把打包内置的默认数据（角色/世界/预设/示例）释放到 exe 旁，仅补缺失不覆盖"""
    bundled = getattr(sys, "_MEIPASS", None)
    if not bundled:
        return
    import shutil
    for name in ("prompt_presets", "personas", "worlds", "saves"):
        src = os.path.join(bundled, name)
        if not os.path.isdir(src):
            continue
        dst = os.path.join(base_dir, name)
        os.makedirs(dst, exist_ok=True)
        for fn in os.listdir(src):
            s = os.path.join(src, fn)
            d = os.path.join(dst, fn)
            if os.path.isfile(s) and not os.path.exists(d):
                try:
                    shutil.copy2(s, d)
                except OSError:
                    pass
    cfg_src = os.path.join(bundled, "config.json")
    cfg_dst = os.path.join(base_dir, "config.json")
    if os.path.isfile(cfg_src) and not os.path.exists(cfg_dst):
        try:
            shutil.copy2(cfg_src, cfg_dst)
        except OSError:
            pass


class HtmlApp:
    def __init__(self):
        self.base_dir = BASE_DIR
        _seed_defaults(self.base_dir)
        self.save_dir = os.path.join(self.base_dir, "saves")
        self.world_dir = os.path.join(self.base_dir, "worlds")
        self.preset_dir = os.path.join(self.base_dir, "prompt_presets")
        self.codex_dir = os.path.join(self.base_dir, "codex")
        self.config_file = os.path.join(self.base_dir, "config.json")
        for d in (self.save_dir, self.world_dir, self.preset_dir, self.codex_dir):
            os.makedirs(d, exist_ok=True)

        self.config = {}
        try:
            with open(self.config_file, "r", encoding="utf-8") as f:
                self.config = json.load(f)
        except Exception:
            pass
        # 架构级水印：config.json 注入来源指纹（一次性，不覆盖已有）
        try:
            import dick_mark
            if not self.config.get("dick_mark"):
                self.config["dick_mark"] = dick_mark.mark_dict({"finger": dick_mark.finger("config")})
                save_guard.atomic_write_json(self.config_file, self.config)
        except Exception:
            pass

        lang = self.config.get("language") or ("en" if i18n.detect_english_system() else "zh")
        i18n.set_lang(lang)
        self.language = lang

        self.core = ChatCore()
        self.core.humanize = bool(self.config.get("humanize", True))
        self._last_user = ""          # 最近一次玩家输入（机制事件判定用）
        self.provider_id = self.config.get("provider") or "deepseek"
        self.api_keys = dict(self.config.get("api_keys") or {})
        key = self.api_keys.get(self.provider_id) or self.config.get("api_key")
        if key:
            self.core.set_api_key(key)
            if self.config.get("base_url"):
                self.core.set_base_url(self.config["base_url"])
        # LLM 通道代理（保证被墙厂商可用）
        if self.config.get("proxy"):
            self.core.set_proxy(self.config["proxy"])
        # 内置代理通道：直连失败自动走中转（默认内置地址，可改）
        self.core.set_relay(self.config.get("relay_url") or BUILTIN_RELAY)
        # 指令模板：停止序列（设置覆盖 > 预设 > 模型家族默认）
        self.core.set_stop_sequences(self._effective_stop())
        # 采样参数（温度/top_p，None=模型默认）
        self.core.set_sampling(self.config.get("temperature"), self.config.get("top_p"))
        if self.config.get("model"):
            self.core.set_model(self.config["model"])
        self.core.set_context_budget(int(self.config.get("context_budget", 0) or 0))
        self.core.set_rolling_summary_enabled(bool(self.config.get("rolling_summary", True)))

        self.plugin_manager = PluginManager(self.core, config_file=self.config_file)
        self.plugin_manager.load_plugins()
        for p in self.plugin_manager.get_all_plugins():
            states = self.config.get("plugin_states", {})
            if p.name in states:
                p.enabled = bool(states[p.name])

        # 前端状态
        self.messages = []          # [{"seq":int,"kind":"user|ai|sys","speaker":str,"content":str,"image":str|None,"node_id":str|None}]
        self.sys_msgs = []          # 系统横幅（欢迎语等，树重建时保留）
        self.node_images = {}       # node_id -> dataURL（图片消息，仅内存，不落盘）
        self.node_seq_cache = {}    # (kind, key) -> seq，重建时保持消息 seq 稳定
        self.seq = 0
        self.render_epoch = 0       # 树重建时 +1，前端据此整刷消息列表
        self.streaming = ""
        self.busy = False
        self.total_tokens = 0
        self.preset_name = ""
        self.persona = None         # {"name","background","notes",...}
        self.auto_turn = bool(self.config.get("auto_turn", False))
        self.font_size = int(self.config.get("ui_font", 0) or 0)
        self.loaded_document = None  # {"name","chars"}
        self._lock = threading.Lock()
        self.ollama_online = False   # 本地 Ollama 是否可用（后台探测）
        self.codex_auto_open = None  # 双击 .codex 文件启动时，待自动打开的包名
        self.codex_volume = 100      # CODEX 播放器音量 0-100
        threading.Thread(target=self._detect_ollama, daemon=True).start()

        self._load_presets()
        self._load_roles()
        self._load_worlds()
        self.current_world = self.config.get("current_world") or ""
        self.core.world_max_entries = int(self.config.get("world_max_entries") or 3)
        self._append_sys(i18n.t("welcome1", "👋 欢迎使用 Direct-Interface Cork-bore Kit v2.0！"))
        self._append_sys(i18n.t("welcome2", "多角色模式：输入 @角色名 内容 来指定说话者。"))
        # 启动即恢复上次会话：激活上次选中的角色（机制卡/恋爱条/战斗随启动初始化，
        # 与 Android 端一致——否则恋爱条要手动保存一遍角色才出现）
        if self.selected_roles:
            try:
                self._activate_core(reload_tree=True)
            except Exception as e:
                print(f"[HtmlApp] 启动恢复选中角色失败: {e}")
        self._rebuild_messages()
        # 存档守护：后台扫描一次，坏档自动修复/从备份恢复
        save_guard.sweep_async(self.base_dir)

    # ---------- 数据 ----------
    def _load_presets(self):
        self.presets = [{"name": "默认"}]
        try:
            for fn in sorted(os.listdir(self.preset_dir)):
                if fn.endswith(".json"):
                    with open(os.path.join(self.preset_dir, fn), "r", encoding="utf-8") as f:
                        data = json.load(f)
                    if isinstance(data, dict) and data.get("name"):
                        self.presets.append(data)
        except Exception:
            pass
        self.preset_name = self.config.get("prompt_preset", "") or ""

    def _load_roles(self):
        self.roles = []
        try:
            for fn in sorted(os.listdir(self.save_dir)):
                if not fn.endswith(".json"):
                    continue
                try:
                    path = os.path.join(self.save_dir, fn)
                    data = save_guard.guard_loaded(path, kind="role")
                    if data is None:
                        continue
                    if isinstance(data, dict) and data.get("system_prompt"):
                        fields = {k: (data.get(k) or "") for k, _ in ROLE_FIELDS}
                        legacy = data.get("legacy") or ""
                        prompt = assemble_role_prompt(data.get("name") or fn[:-5], fields, legacy) \
                            or data.get("system_prompt", "")
                        self.roles.append({"name": data.get("name") or fn[:-5],
                                           "file": fn, "prompt": prompt, "data": data,
                                           "fields": fields, "legacy": legacy,
                                           "unlocked": bool(data.get("unlocked", False))})
                except Exception:
                    continue
        except Exception:
            pass
        # 恢复上次选中的角色（与 Android 端一致：启动即恢复，恋爱条/机制卡初始就绪，
        # 无需再手动保存一遍才出现）。config.selected_roles（新版数组）或 last_role（旧版单值）。
        self.selected_roles = []
        try:
            saved = self.config.get("selected_roles") or []
            if not isinstance(saved, list):
                saved = []
            if not saved and self.config.get("last_role"):
                lr = str(self.config.get("last_role") or "").strip()
                if lr:
                    saved = [lr]
            self.selected_roles = [n for n in saved
                                   if any(r["name"] == n for r in self.roles)][:8]
        except Exception:
            self.selected_roles = []

    def _detect_ollama(self):
        """后台探测本地 Ollama（localhost / 127.0.0.1），结果进 api_state"""
        try:
            import requests
            for url in ("http://localhost:11434/v1/models",
                        "http://127.0.0.1:11434/v1/models"):
                try:
                    r = requests.get(url, timeout=3)
                    if r.status_code < 500:
                        self.ollama_online = True
                        print("[Ollama] ✅ 本地 Ollama 已检测到")
                        return
                except Exception:
                    continue
        except Exception:
            pass
        self.ollama_online = False
        print("[Ollama] 未检测到本地服务（需要时先安装并启动 Ollama）")

    def _vision_describe(self, image_b64, mime):
        """走免费视觉链描述图片；配置了代理时一并走代理"""
        import requests
        proxy = self.config.get("proxy") or None
        proxies = {"http": proxy, "https": proxy} if proxy else None
        return _vision_describe(image_b64, mime, proxies=proxies)

    def _load_worlds(self):
        self.worlds = []
        try:
            for fn in sorted(os.listdir(self.world_dir)):
                if not fn.endswith(".json"):
                    continue
                try:
                    data = save_guard.guard_loaded(os.path.join(self.world_dir, fn), kind="world")
                    if isinstance(data, dict) and data.get("name"):
                        self.worlds.append(data)
                except Exception:
                    continue
        except Exception:
            pass
        self.selected_worlds = []

    def _append_sys(self, content, kind="sys", speaker="系统"):
        """树外横幅（系统提示/命令回显），随重建一直显示"""
        with self._lock:
            self.sys_msgs.append({"kind": kind, "speaker": speaker, "content": content})
            if len(self.sys_msgs) > 200:
                self.sys_msgs = self.sys_msgs[-200:]
        self._rebuild_messages()

    def _chain_nodes(self):
        """从 core 树里取当前链（跳过 system 节点），返回 MessageNode 列表（root→leaf）"""
        tree = self.core.tree
        leaf = tree.current_leaf_id
        if not leaf or leaf not in tree.nodes:
            return []
        chain = []
        node = tree.nodes.get(leaf)
        while node is not None:
            chain.append(node)
            if node.parent_id and node.parent_id in tree.nodes:
                node = tree.nodes.get(node.parent_id)
            else:
                break
        chain.reverse()
        return [n for n in chain if n.role != "system"]

    def _rebuild_messages(self):
        """按 当前聊天树链 + 系统横幅 重建展示消息列表（epoch +1 通知前端整刷）。
        消息 seq 按 (类型,键) 缓存：重建不重排，前端旧 seq 依然有效。"""
        with self._lock:
            def nxt(key):
                if key in self.node_seq_cache:
                    return self.node_seq_cache[key]
                self.seq += 1
                self.node_seq_cache[key] = self.seq
                return self.seq

            msgs = []
            for i, s in enumerate(self.sys_msgs):
                if isinstance(s, str):
                    s = {"kind": "sys", "speaker": "系统", "content": s}
                msgs.append({"seq": nxt(("sys", i)), "kind": s.get("kind", "sys"),
                             "speaker": s.get("speaker", "系统"),
                             "content": s.get("content", ""), "image": None, "node_id": None})
            for n in self._chain_nodes():
                seq = nxt(("node", n.id))
                meta = n.metadata or {}
                kind = "user" if n.role == "user" else "ai"
                speaker = meta.get("speaker")
                if kind == "user":
                    # 用户消息显示玩家卡名字（与头像文件名一致），无名字时显示「你」
                    speaker = (self.persona or {}).get("name") or "你"
                elif not speaker:
                    # 单角色：AI 消息显示当前激活角色名；群聊/无角色时显示 AI
                    speaker = self.selected_roles[0] if len(self.selected_roles) == 1 else "AI"
                msgs.append({"seq": seq, "kind": kind, "speaker": speaker,
                             "content": n.content, "image": self.node_images.get(n.id),
                             "node_id": n.id})
            # 滑条信息：AI 消息的父用户节点有多个 assistant 子节点时给箭头
            for i, m in enumerate(msgs):
                if m["kind"] != "ai" or not m["node_id"]:
                    continue
                node = self.core.tree.nodes.get(m["node_id"])
                if node and node.parent_id and node.parent_id in self.core.tree.nodes:
                    sibs = [cid for cid in self.core.tree.nodes[node.parent_id].children_ids
                            if self.core.tree.nodes.get(cid) and self.core.tree.nodes[cid].role == "assistant"]
                    if len(sibs) > 1:
                        m["swipes"] = {"total": len(sibs), "index": sibs.index(m["node_id"])}
            self.messages = msgs
            self.render_epoch += 1

    def _save_tree(self):
        """把当前聊天树写回第一个选中角色的文件（保留 name/system_prompt/card_data）"""
        if not self.selected_roles:
            return
        tree_data = self.core.get_all_nodes_data()
        # 空树（只有系统节点）不覆盖已有历史，防止切换/误操作清档
        if len(tree_data.get("nodes", {})) <= 1:
            return
        for r in self.roles:
            if r["name"] == self.selected_roles[0]:
                data = dict(r.get("data") or {})
                data["name"] = r["name"]
                data["system_prompt"] = r["prompt"]
                data["history_tree"] = self.core.get_all_nodes_data()
                try:
                    # 存档守护：覆盖前留底 + 原子写入（写一半崩溃也不会截断存档）
                    save_guard.backup_file(os.path.join(self.save_dir, r["file"]))
                    save_guard.atomic_write_json(os.path.join(self.save_dir, r["file"]), data)
                except Exception:
                    pass
                # 同步内存快照：勾选/取消勾选切换时 _activate_core 从
                # r["data"]["history_tree"] 取记录，不同步会导致同一会话内
                # 重新勾选只恢复到旧快照（新建角色的记录甚至取不回）
                r["data"] = data
                break

    # ---------- 系统提示 ----------
    def _build_system_prompt(self):
        sb = []
        preset = None
        for p in self.presets:
            if p.get("name") == self.preset_name:
                preset = p
                break
        if preset and preset.get("system_prefix"):
            sb.append(preset["system_prefix"])
        chosen = [r for r in self.roles if r["name"] in self.selected_roles]
        for r in chosen:
            sb.append(r["prompt"])
        if len(chosen) > 1:
            sb.append("当前是群聊，参与角色：" + "、".join(r["name"] for r in chosen) +
                      "。每次由一位角色发言，用 [角色名]: 开头标注。用户 @角色名 表示指定该角色回复。")
        ws = [w for w in self.worlds if w.get("name") in self.selected_worlds]
        if ws:
            sb.append("【世界设定】")
            for w in ws:
                sb.append(str(w.get("name")) + "：" + str(w.get("description", "")))
        if self.persona:
            sb.append("【玩家角色卡】" + json.dumps(self.persona, ensure_ascii=False))
        if preset and preset.get("rules"):
            sb.append(preset["rules"])
        injection = self.plugin_manager.contextInjection() if hasattr(self.plugin_manager, "contextInjection") else ""
        if injection:
            sb.append(injection)
        if preset and preset.get("system_suffix"):
            sb.append(preset["system_suffix"])
        return "\n\n".join(x for x in sb if x and str(x).strip())

    # ---------- Galgame 选项（插件联动） ----------
    def _effective_stop(self):
        """生效的停止序列：设置覆盖 > 预设 stop_sequences > 当前模型家族默认"""
        ov = self.config.get("stop_sequences") or []
        if ov:
            return [str(x).strip() for x in ov if str(x).strip()]
        preset = getattr(self.core, "prompt_preset", None) or {}
        ps = preset.get("stop_sequences") or []
        if ps:
            return [str(x).strip() for x in ps if str(x).strip()]
        prov = next((p for p in PROVIDERS if p["id"] == self.provider_id), None)
        return list((prov or {}).get("stop", []) or [])

    def api_set_stop(self, stop_text):
        """设置停止序列（逗号/换行分隔；空 = 用预设/模型默认）"""
        parts = (stop_text or "").replace("，", ",").replace("\n", ",").split(",")
        self.config["stop_sequences"] = [x.strip() for x in parts if x.strip()]
        self.core.set_stop_sequences(self._effective_stop())
        self._save_config()
        return {"ok": True, "stop": list(self.core.stop_sequences or [])}

    def api_set_sampling(self, temperature, top_p):
        """设置采样参数（温度/top_p；空 = 模型默认）"""
        self.config["temperature"] = temperature
        self.config["top_p"] = top_p
        self.core.set_sampling(temperature, top_p)
        self._save_config()
        return {"ok": True,
                "temperature": self.core.temperature,
                "top_p": self.core.top_p}

    def _choices_plugin(self):
        try:
            if self.plugin_manager:
                return self.plugin_manager.get_plugin("Galgame 选项")
        except Exception:
            pass
        return None

    def _choices_state(self):
        p = self._choices_plugin()
        if not p or not p.enabled:
            return {"items": [], "loading": False, "error": ""}
        try:
            return {"items": list(getattr(p, "choices", None) or []),
                    "loading": bool(getattr(p, "choices_loading", False)),
                    "error": str(getattr(p, "choices_error", "") or "")}
        except Exception:
            return {"items": [], "loading": False, "error": ""}

    def _clear_choices(self):
        p = self._choices_plugin()
        if p:
            try:
                p.clear_choices()
            except Exception:
                pass

    # ---------- js_api ----------
    def api_state(self):
        with self._lock:
            return {
                "lang": self.language,
                "roles": [r["name"] for r in self.roles],
                "role_unlocked": {r["name"]: bool(r.get("unlocked", False)) for r in self.roles},
                "worlds": [w.get("name", "") for w in self.worlds],
                "presets": [p.get("name", "默认") for p in self.presets],
                "selected_roles": list(self.selected_roles),
                "selected_worlds": list(self.selected_worlds),
                "current_world": self.current_world,
                "preset": self.preset_name,
                "persona": self.persona,
                "has_key": bool(self.api_keys.get(self.provider_id) or self.config.get("api_key")),
                "model": self.config.get("model", "deepseek-v4-flash"),
                "base_url": self.config.get("base_url", "https://api.deepseek.com"),
                "providers": PROVIDERS,
                "provider": self.provider_id,
                "proxy": self.config.get("proxy", ""),
                "relay_url": self.config.get("relay_url") or BUILTIN_RELAY,
                "relay_on": bool(getattr(self.core, "relay_on", False)),
                "stop_input": list(self.config.get("stop_sequences") or []),
                "stop_sequences": list(getattr(self.core, "stop_sequences", None) or []),
                "temperature": self.config.get("temperature"),
                "top_p": self.config.get("top_p"),
                "dev_mode": bool(self.config.get("dev_mode", False)),
                "humanize": bool(getattr(self.core, "humanize", True)),
                "welcome_shown": bool(self.config.get("welcome_shown", False)),
                "ollama_online": bool(getattr(self, "ollama_online", False)),
                "api_keys": self.api_keys,
                "auto_turn": self.auto_turn,
                "font": self.font_size,
                "budget": int(self.config.get("context_budget", 0) or 0),
                "document": self.loaded_document,
                "plugins": [{"name": p.name, "version": p.version, "description": p.description,
                             "enabled": bool(p.enabled),
                             "schema": getattr(p, "settings_schema", None) or [],
                             "settings": dict(getattr(p, "settings", {}) or {})}
                            for p in self.plugin_manager.get_all_plugins()],
                "theme": self.config.get("ui_theme", 0),
                "accent": self.config.get("ui_accent", 0),
                "messages": self.messages,
                "streaming": self.streaming,
                "busy": self.busy,
                "tokens": self.total_tokens,
                "epoch": self.render_epoch,
                "choices": self._choices_state(),
                "mechanism": {"config": getattr(self.core, "_mech_config", None),
                              "state": self.core.mechanism_snapshot()},
                "battle": self.core.battle_ui_state(),
                "codex_auto_open": self.codex_auto_open,
            }

    def api_poll(self, last_seq):
        with self._lock:
            items = [m for m in self.messages if m["seq"] > last_seq]
            return {"items": items, "streaming": self.streaming, "busy": self.busy,
                    "tokens": self.total_tokens, "epoch": self.render_epoch,
                    "choices": self._choices_state(),
                    "mechanism": {"config": getattr(self.core, "_mech_config", None),
                                  "state": self.core.mechanism_snapshot()},
                    "battle": self.core.battle_ui_state()}

    def _roll_option(self, item):
        """点选项时按配置概率表 ROLL（百分比可配：机制卡 mechanics.roll，默认
        暴击10/稀有4/大失败2/天选0.1(千分之一)/坍缩0.001(十万分之一)；概率不公布）。
        返回 (effect, kind, note)"""
        import random as _r
        cfg = getattr(self.core, "_mech_config", None) or {}
        roll = cfg.get("roll") if isinstance(cfg, dict) else None
        if not isinstance(roll, dict):
            roll = {}
        def _p(key, dflt):
            try:
                return float(roll.get(key, dflt))
            except (TypeError, ValueError):
                return dflt
        # 天选 = 千分之一（兼容旧键 legend：旧卡没配 chosen 时沿用其传说概率）
        chosen_p = _p("chosen", _p("legend", 0.1))
        collapse_p = _p("collapse", 0.001)
        fail_p = _p("fail", 2)
        rare_p = _p("rare", 4)
        crit_p = _p("crit", 10)
        chosen_p = max(0.0, min(100.0, chosen_p)) / 100.0
        collapse_p = max(0.0, min(100.0, collapse_p)) / 100.0
        fail_p = max(0.0, min(100.0, fail_p)) / 100.0
        rare_p = max(0.0, min(100.0, rare_p)) / 100.0
        crit_p = max(0.0, min(100.0, crit_p)) / 100.0
        table = [
            (collapse_p, "collapse", "🌌 坍缩：十万分之一的奇迹坍缩成现实！"),
            (collapse_p + chosen_p, "chosen", "🌟 天选：千分之一的天命眷顾被触发了！"),
            (collapse_p + chosen_p + fail_p, "fail", "💥 结果出了岔子！"),
            (collapse_p + chosen_p + fail_p + rare_p, "rare", "✨ 命运的眷顾：触发了稀有事件！"),
            (collapse_p + chosen_p + fail_p + rare_p + crit_p, "crit", "✨ 效果暴击！"),
            (1.0, "normal", ""),
        ]
        r = _r.random()
        kind, note = "normal", ""
        for limit, k, n in table:
            if r < limit:
                kind, note = k, n
                break
        effect = {}
        if isinstance(item, dict):
            aff = item.get("aff")
            st = item.get("st")
            if isinstance(st, dict):
                effect["st"] = st
            if aff is not None:
                try:
                    aff = int(aff)
                except (TypeError, ValueError):
                    aff = None
            if aff is not None:
                if kind == "crit":
                    effect["aff"] = aff * 2
                elif kind == "fail":
                    effect["aff"] = -aff
                else:
                    effect["aff"] = aff
            elif kind in ("crit", "fail"):
                # 无机制效果的选项：暴击/失败仅提示，不空转
                note = ""
        return effect, kind, note

    def api_pick_choice(self, text):
        """Galgame 选项：点击某个选项 = 隐藏 ROLL 出结果 + 以该行动作为玩家输入发言"""
        text = (text or "").strip()
        if not text:
            return {"ok": False, "err": "empty"}
        if self.busy:
            return {"ok": False, "err": "busy"}
        # 选项效果（好感/状态）：从插件当前选项里按文本匹配
        p = self._choices_plugin()
        item = None
        if p:
            try:
                for it in (getattr(p, "choices", None) or []):
                    if isinstance(it, dict) and it.get("text") == text:
                        item = it
                        break
            except Exception:
                pass
        # 隐藏 ROLL：按内置概率表出结果
        effect, kind, note = self._roll_option(item or {"text": text})
        self.core.apply_mechanism_effect(effect)
        if kind == "rare":
            self.core.pending_event = {"id": "_roll_rare", "name": "命运的眷顾",
                                       "prompt": "（稀有事件）这段剧情出现了意想不到的转折，"
                                                 "请自然地演出一个令人惊喜的展开。"}
            if note:
                self._append_sys(note)
        elif kind == "chosen":
            self.core.pending_event = {"id": "_roll_chosen", "name": "天选",
                                       "prompt": "（天选事件）千分之一的天命眷顾发生了！"
                                                 "请演出一个不可思议的、足以载入史册的剧情转折。"}
            if note:
                self._append_sys(note)
        elif kind == "collapse":
            self.core.pending_event = {"id": "_roll_collapse", "name": "坍缩",
                                       "prompt": "（坍缩事件）十万分之一的奇迹坍缩成现实！"
                                                 "战斗数值全部坍缩为 2000。"
                                                 "请演出一个撼动世界观的、堪称神话的剧情展开。"}
            try:
                self.core.collapse_battle_values()
            except Exception:
                pass
            if note:
                self._append_sys(note)
        elif kind in ("crit", "fail") and note:
            self._append_sys(note)
        self._clear_choices()
        self._send_text(text, None, None)
        return {"ok": True}

    def api_cyoa(self):
        """Galgame 选项：手动生成一组剧情选项（不写进聊天记录）"""
        p = self._choices_plugin()
        if not p or not p.enabled:
            return {"ok": False, "err": "插件未启用（⚙️ 设置 → 插件 → Galgame 选项）"}
        ok, msg = p.manual_generate()
        return {"ok": ok, "msg": msg}

    def api_battle_move(self, move_id):
        """战斗：玩家出招 → 引擎结算伤害/消耗 → 结算文本作为系统横幅 + 行动作为玩家消息发出（AI 演出）"""
        move_id = (move_id or "").strip()
        if not move_id:
            return {"ok": False, "err": "empty"}
        if self.busy:
            return {"ok": False, "err": "busy"}
        cfg = self.core._battle_config()
        if not cfg:
            return {"ok": False, "err": "当前角色未启用战斗系统"}
        move_name = move_id
        for m in (cfg.get("moves") or []):
            if isinstance(m, dict) and str(m.get("id", "")).strip() == move_id:
                move_name = str(m.get("name") or m.get("id"))
                break
        txt, is_legend = self.core.resolve_battle_move(move_id)
        if txt is None:
            return {"ok": False, "err": "招式不存在"}
        # 天选之人：0.00001% 战斗奇迹
        if is_legend:
            self.core.pending_event = {"id": "_battle_legend", "name": "天选之人",
                                       "prompt": "（传说事件）战斗中发生了十万分之一的奇迹！"
                                                 "请演出一个足以载入史册的惊天转折。"}
            self._append_sys("🌟 天选之人：十万分之一的战斗奇迹被触发了！")
        if txt.startswith("⚠️"):
            self._append_sys(txt)
            return {"ok": True, "msg": txt}
        self._append_sys(txt)
        # 结算结果注入下一轮请求（AI 知道伤害数字，演出受击反应）
        self.core.pending_event = {"id": "_battle_result", "name": "战斗结算",
                                   "prompt": f"（战斗结算）{txt} 请以角色口吻演出受击反应与战况。"}
        self._send_text(f"使用 {move_name}", None, None)
        return {"ok": True, "msg": txt}

    def _expand_macros(self, text):
        """Quick Reply 宏展开（酒馆 {{player}}/{{char}} 的对等物）：
        {player} 玩家卡名 · {char} 当前角色 · {world} 当前世界 · {random:a|b|c} 随机取一
        未知宏原样保留。"""
        if not text or "{" not in text:
            return text
        import random as _r
        player = (self.persona or {}).get("name") or "你"
        char = self.selected_roles[0] if self.selected_roles else "AI"
        world = self.current_world or ""

        def rep(m):
            key = m.group(1).strip()
            if key == "player":
                return player
            if key == "char":
                return char
            if key == "world":
                return world
            if key.startswith("random:"):
                opts = [x for x in key[len("random:"):].split("|") if x]
                return _r.choice(opts) if opts else ""
            return m.group(0)

        import re as _re
        return _re.sub(r"\{([^{}]+)\}", rep, text)

    def api_quick_replies(self):
        """Quick Reply 列表：全局 quick_replies.json + 当前激活卡片的 card_quick_replies"""
        out = []
        try:
            with open(os.path.join(self.base_dir, "quick_replies.json"), encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                out.extend({"label": str(x.get("label") or ""), "text": str(x.get("text") or "")}
                           for x in data if isinstance(x, dict) and x.get("label"))
        except Exception:
            pass
        # 当前卡片自带快捷回复（高级设置）排在最前
        for name in self.selected_roles:
            for r in self.roles:
                if r["name"] == name:
                    adv = (r.get("data") or {}).get("advanced") or {}
                    for q in adv.get("card_quick_replies") or []:
                        if q.get("label"):
                            out.insert(0, {"label": str(q["label"]), "text": str(q.get("text") or "")})
                    break
        return out

    # ---------- 正则管道（ST 风格：清洗/格式化；存储前应用，树状天然一致） ----------
    def _regex_rules_all(self, scope):
        """生效规则：全局 regex_rules.json + 当前角色卡 regex_rules（角色卡优先）"""
        rules = []
        try:
            with open(os.path.join(self.base_dir, "regex_rules.json"), encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                rules.extend(data)
        except Exception:
            pass
        for name in self.selected_roles:
            for r in self.roles:
                if r["name"] == name:
                    adv = (r.get("data") or {}).get("advanced") or {}
                    rr = adv.get("regex_rules") or []
                    if isinstance(rr, list):
                        rules = list(rr) + rules
                    break
        # 玩家卡规则（同规格待遇）：对玩家输入生效，优先于角色卡
        p = self.persona or {}
        if isinstance(p, str):
            try:
                p = json.loads(p)
            except Exception:
                p = {}
        if isinstance(p, dict):
            prr = (p.get("advanced") or {}).get("regex_rules") or []
            if isinstance(prr, list):
                rules = list(prr) + rules
        out = []
        for x in rules:
            if not isinstance(x, dict) or not x.get("enabled", True):
                continue
            if not str(x.get("pattern") or "").strip():
                continue
            if (x.get("scope") or "both") not in (scope, "both"):
                continue
            out.append(x)
        return out

    def _apply_regex_pipeline(self, text, scope):
        """应用正则管道（存储前调用 → 树里存转换后文本，回溯/分支天然一致）"""
        if not text:
            return text
        for rule in self._regex_rules_all(scope):
            try:
                pattern = str(rule.get("pattern") or "")
                repl = str(rule.get("replace") or "")
                text = re.sub(pattern, repl, text)
            except Exception:
                continue
        return text

    def api_get_regex_rules(self):
        """全局正则规则（编辑预填）"""
        rules = []
        try:
            with open(os.path.join(self.base_dir, "regex_rules.json"), encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                rules = data
        except Exception:
            pass
        return {"rules": rules}

    def api_set_regex_rules(self, rules_json):
        """保存全局正则规则"""
        try:
            parsed = json.loads(rules_json) if isinstance(rules_json, str) else rules_json
            if not isinstance(parsed, list):
                return {"ok": False, "err": "格式错误"}
            cleaned = []
            for x in parsed:
                if not isinstance(x, dict) or not str(x.get("pattern") or "").strip():
                    continue
                scope = str(x.get("scope") or "both")
                if scope not in ("ai", "user", "both"):
                    scope = "both"
                cleaned.append({
                    "id": str(x.get("id") or "")[:30],
                    "name": str(x.get("name") or x.get("id") or "")[:30],
                    "pattern": str(x.get("pattern") or "").strip(),
                    "replace": str(x.get("replace") or ""),
                    "scope": scope,
                    "enabled": bool(x.get("enabled", True)),
                })
            with open(os.path.join(self.base_dir, "regex_rules.json"), "w", encoding="utf-8") as f:
                json.dump(cleaned, f, ensure_ascii=False, indent=2)
            return {"ok": True, "count": len(cleaned)}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_set_dev_mode(self, enabled):
        """开发者模式开关（解锁角色卡高级设置：内置游戏等）"""
        self.config["dev_mode"] = bool(enabled)
        self._save_config()
        return {"ok": True, "dev_mode": bool(enabled)}

    def api_set_humanize(self, enabled):
        """去 AI 味开关（默认开）：注入人性化对话规则"""
        self.config["humanize"] = bool(enabled)
        self.core.humanize = bool(enabled)
        self._save_config()
        return {"ok": True, "humanize": bool(enabled)}

    def api_dismiss_welcome(self):
        """标记首次引导已看完（不再弹引导框）"""
        self.config["welcome_shown"] = True
        self._save_config()
        return {"ok": True}

    def api_send(self, text, image_b64=None, mime=None):
        text = (text or "").strip()
        if self.busy:
            return {"ok": False, "err": "busy"}
        if not text and not image_b64:
            return {"ok": False, "err": "empty"}
        if text.startswith("/"):
            travel = self._handle_travel(text)
            if travel:
                self._append_sys(text, kind="user", speaker="你")
                self._append_sys(travel)
                return {"ok": True}
            if text.strip().lower().startswith("/mettertools"):
                # METTERTOOLS：按上限百分比一键填好感（罪恶都市梗；/mettertools 90 = 填到 90%）
                self._append_sys(text, kind="user", speaker="你")
                parts = text.split()
                pct = int(parts[1]) if len(parts) > 1 and parts[1].isdigit() else 100
                val = self.core.set_affection_percent(pct)
                if val is None:
                    self._append_sys("⚠️ 当前角色未启用好感度（机制卡 → ❤ 好感度）")
                else:
                    self._append_sys(f"✨ METTERTOOLS！好感度已填至 {pct}% → {val}。")
                return {"ok": True}
            if text.strip() == "/read" or text.startswith("/read "):
                self._append_sys(text, kind="user", speaker="你")
                arg = text[5:].strip() if text.startswith("/read ") else None
                r = self.api_read_document(arg)
                if r.get("ok"):
                    self._append_sys("✅ 已注入文档上下文（可用 /readclear 清除）")
                elif r.get("err") != "cancelled":
                    self._append_sys("⚠️ " + r.get("err", "读取失败"))
                return {"ok": True}
            if text.strip() == "/readclear":
                self._append_sys(text, kind="user", speaker="你")
                self.api_clear_document()
                return {"ok": True}
            result = self.plugin_manager.handle_command(text)
            self._append_sys(text, kind="user", speaker="你")
            self._append_sys(result or i18n.t("unknown_cmd", "未知命令"))
            return {"ok": True}
        # Quick Reply 宏展开：{player} {char} {world} {random:a|b|c}
        text = self._expand_macros(text)
        processed = self.plugin_manager.onMessageSend(text) if hasattr(self.plugin_manager, "onMessageSend") else text
        if processed is None:
            return {"ok": False, "err": "blocked"}
        # 翻译隐藏：显示/存储用原文，发给 AI 用译文（中字日配不露痕迹）
        display_text, send_text = self._split_hidden(processed)
        # 玩家发送新消息（打字或点选项之外的自由输入）后，旧剧情选项作废
        self._clear_choices()
        image = None
        if image_b64:
            image = "data:" + (mime or "image/png") + ";base64," + image_b64
        # 传图补丁：先视觉描述（后台线程）
        if image_b64:
            self.busy = True
            speaker = (self.persona or {}).get("name") or "你"
            node_id = self.core.add_user_message(display_text)  # 树里存原文（翻译隐藏）
            try:
                self.core.tree.nodes[node_id].metadata["speaker"] = speaker
            except Exception:
                pass
            self.node_images[node_id] = image
            self._rebuild_messages()
            threading.Thread(target=self._vision_then_send,
                             args=(node_id, send_text, image_b64, mime), daemon=True).start()  # 发译文给 AI
            return {"ok": True}
        self._send_text(display_text, send_text, None)

        return {"ok": True}

    def _vision_then_send(self, node_id, text, image_b64, mime):
        try:
            desc = self._vision_describe(image_b64, mime or "image/png")
        except Exception as e:
            desc = None
            print(f"[视觉] 描述失败: {e}")
        node = self.core.tree.nodes.get(node_id)
        if desc:
            tree_content = "用户发送了一张图片。视觉模型对图片的描述：\n" + desc + "\n\n用户输入：" + text
            if node:
                node.content = tree_content
            self._rebuild_messages()
            self._start_fetch(node_id)
        else:
            self.busy = False
            self._append_sys(i18n.t("vision_fail", "⚠️ 图片识别失败（免费视觉链被限流或网络问题），请稍后重试"))

    def _handle_travel(self, text):
        """/穿越 [世界名]：列出可穿越世界或直接穿越；返回提示文本或 None（不是穿越命令）"""
        parts = text.split(None, 1)
        if (parts[0] if parts else "").strip().lower() not in ("/穿越", "/travel"):
            return None
        target = parts[1].strip() if len(parts) > 1 else ""
        if not self.selected_worlds:
            return "⚠️ 还没有勾选世界卡（左侧世界列表勾选后即可穿梭）"
        if not target:
            lines = ["🌍 可穿越的世界："]
            for i, n in enumerate(self.selected_worlds, 1):
                mark = " ★当前所在" if n == self.current_world else ""
                lines.append(str(i) + ". " + n + mark)
            lines.append("用法：/穿越 世界名")
            return chr(10).join(lines)
        hit = next((n for n in self.selected_worlds if n == target or target in n), None)
        if not hit:
            return "⚠️ 没有找到世界「" + target + "」，可用 /穿越 查看列表"
        self.api_travel(hit)
        return "🚀 已穿越到「" + hit + "」"

    def _start_fetch(self, user_node_id):
        """在指定用户节点下生成回复（候选/滑条/分支共用入口）"""
        self.busy = True
        self.streaming = ""
        self.core.generate_candidate(
            user_node_id,
            on_stream=lambda full: setattr(self, "streaming",
                                           self.core.strip_mechanism_tags(full, apply=False)),
            on_response=self._on_response,
            on_error=self._on_error,
        )

    def _split_hidden(self, processed):
        """翻译隐藏标记解析 → (显示文本, 发送文本)。无标记则两者相同。
        约定：jp_patch 自动日译返回 "\u200b<原文>\u200b<译文>"。"""
        if not processed:
            return processed, processed
        mark = "\u200b"
        if isinstance(processed, str) and processed.startswith(mark):
            parts = processed.split(mark)
            if len(parts) >= 3:
                return parts[1], mark.join(parts[2:])
        return processed, processed

    def _send_text(self, display_text, hidden_send, image):
        # 显示/存储用原文；若翻译隐藏（hidden_send 非空且不同），译文存 metadata["ja_input"]，
        # _fetch_response 发 AI 时优先用译文（聊天永远看不到日文）
        content = self._apply_regex_pipeline(display_text or "", "user")
        self._last_user = display_text or ""
        speaker = (self.persona or {}).get("name") or "你"
        node_id = self.core.add_user_message(content)
        try:
            self.core.tree.nodes[node_id].metadata["speaker"] = speaker
            if hidden_send and hidden_send != display_text:
                self.core.tree.nodes[node_id].metadata["ja_input"] = hidden_send
            if self.core.mechanism_state is not None:
                self.core.tree.nodes[node_id].metadata["ms"] = self.core.mechanism_snapshot()
        except Exception:
            pass
        if image:
            self.node_images[node_id] = image
        self._start_fetch(node_id)
        self._rebuild_messages()

    def _on_response(self, ai_reply, usage):
        if usage:
            self.total_tokens += int(getattr(usage, "total_tokens", 0) or 0)
        # 中字日配：提取 [ja]...[/ja] 隐藏配音句（存节点 metadata，正文剥离）
        ja_text = ""
        try:
            m = re.search(r"\[ja\]([\s\S]*?)\[/ja\]", ai_reply or "")
            if m:
                ja_text = m.group(1).strip()
                ai_reply = re.sub(r"\[ja\][\s\S]*?\[/ja\]", "", ai_reply)
        except Exception:
            pass
        # 机制卡：解析 [aff:+N]/[键:值] 标签 → 应用状态 → 更新节点内容 → 检查事件
        try:
            clean = self.core.strip_mechanism_tags(ai_reply or "", apply=True)
            # 正则管道（ai 作用域）：标签剥离后、写入树前应用
            clean = self._apply_regex_pipeline(clean, "ai")
            if self.core.mechanism_state is not None:
                leaf = self.core.tree.current_leaf_id
                if leaf and leaf in self.core.tree.nodes:
                    node = self.core.tree.nodes[leaf]
                    if clean != (node.content or ""):
                        node.content = clean
                    ev = self.core.check_mech_events(self._last_user)
                    if ev:
                        self.core.pending_event = ev
                    node.metadata["ms"] = self.core.mechanism_snapshot()
                    if ja_text:
                        node.metadata["ja"] = ja_text
            ai_reply = clean
        except Exception:
            pass
        self._save_tree()
        self._rebuild_messages()
        self.streaming = ""
        self.busy = False
        if self.auto_turn and len(self.selected_roles) > 1:
            self._maybe_auto_turn(2)
        for p in self.plugin_manager.get_all_plugins():
            if p.enabled:
                try:
                    p.on_message_received(getattr(self, "_last_user", ""), ai_reply)
                except Exception:
                    pass

    def _on_error(self, err_msg):
        self.streaming = ""
        self.busy = False
        self._save_tree()
        self._append_sys("❌ " + str(err_msg))

    # ---------- 管理动作 ----------
    def _activate_core(self, reload_tree=True):
        """把选中角色/世界/预设/玩家卡接入 core（重建系统节点 + 首个角色的历史树）"""
        roles_data = []
        for i, name in enumerate(self.selected_roles):
            for r in self.roles:
                if r["name"] == name:
                    rd = {"name": r["name"], "system_prompt": r["prompt"],
                          "unlocked": r.get("unlocked", False),
                          "advanced": (r.get("data") or {}).get("advanced")}
                    if i == 0 and reload_tree and r.get("data") and r["data"].get("history_tree"):
                        rd["history_tree"] = r["data"]["history_tree"]
                    roles_data.append(rd)
        self.core.set_active_roles(roles_data)
        self.core.set_worlds(self._worlds_for_core())
        self.core.set_player_persona(self.persona)
        preset = None
        for p in self.presets:
            if p.get("name") == self.preset_name:
                preset = p
                break
        self.core.set_prompt_preset(preset if preset and preset.get("name") else None)
        # 切换角色 = 新会话：旧剧情选项作废
        self._clear_choices()
        self.core.pending_event = None
        self._rebuild_messages()

    def api_select_roles(self, names_json):
        try:
            names = json.loads(names_json)
            names = [n for n in names if any(r["name"] == n for r in self.roles)]
        except Exception:
            return {"ok": False}
        # 先把当前聊天树写回「旧」的首个选中角色，再切换选择：
        # 否则取消勾选/切换角色时，旧角色的记录会残留在聊天窗口，
        # 甚至误写进新选角色的存档里。
        self._save_tree()
        self.selected_roles = names
        self.config["selected_roles"] = list(names)
        self.config["last_role"] = names[0] if names else ""
        self._save_config()
        # 切换角色 = 新会话：清空系统横幅（欢迎语/命令回显等），界面彻底初始化
        self.sys_msgs = []
        self._activate_core(reload_tree=True)
        return {"ok": True}

    def _worlds_for_core(self):
        """选中世界 → core 数据结构（背景中内嵌世界参数的人类可读渲染）"""
        out = []
        for w in self.worlds:
            if w.get("name") in self.selected_worlds:
                wc = dict(w)
                wc["description"] = _render_world_desc(w)
                out.append(wc)
        return out

    def _sync_worlds(self):
        """把选中世界接入 core；当前世界若不在选中列表则回退到第一个"""
        selected = self._worlds_for_core()
        self.core.set_worlds(selected)
        if self.current_world not in self.selected_worlds:
            self.current_world = self.selected_worlds[0] if self.selected_worlds else ""
        if self.current_world:
            self.core.set_current_world(self.current_world)
        self.config["current_world"] = self.current_world
        self._save_config()

    def api_select_worlds(self, names_json):
        try:
            names = json.loads(names_json)
            self.selected_worlds = [n for n in names if any(w.get("name") == n for w in self.worlds)]
        except Exception:
            pass
        self._sync_worlds()
        return {"ok": True}

    def api_travel(self, name):
        """穿梭：把角色移动到指定世界（世界书条目只跟随当前世界）"""
        name = (name or "").strip()
        if not name or name not in self.selected_worlds:
            return {"ok": False, "err": "请先勾选该世界再穿越"}
        self.current_world = name
        self.core.set_current_world(name)
        self.config["current_world"] = name
        self._save_config()
        return {"ok": True, "current": name}

    def api_set_preset(self, name):
        self.preset_name = name or ""
        self.config["prompt_preset"] = self.preset_name
        self._save_config()
        preset = None
        for p in self.presets:
            if p.get("name") == self.preset_name:
                preset = p
                break
        self.core.set_prompt_preset(preset if preset and preset.get("name") else None)
        # 指令模板：预设变更后重新应用停止序列（预设可带 stop_sequences）
        self.core.set_stop_sequences(self._effective_stop())
        return {"ok": True}

    def api_get_persona(self):
        """返回玩家角色卡结构化字段（供精细化表单预填）"""
        p = self.persona
        if isinstance(p, str):
            try:
                p = json.loads(p)
            except Exception:
                p = None
        p = p or {}
        return {
            "name": str(p.get("name") or ""),
            "legacy": str(p.get("legacy") or ""),
            "appearance": str(p.get("appearance") or ""),
            "personality": str(p.get("personality") or ""),
            "background": str(p.get("background") or ""),
            "speech": str(p.get("speech") or ""),
            "first_mes": str(p.get("first_mes") or ""),
            "mes_example": str(p.get("mes_example") or ""),
            "notes": str(p.get("notes") or ""),
            "advanced": p.get("advanced") if isinstance(p.get("advanced"), dict) else None,
        }

    def api_set_persona(self, persona_json):
        old = (self.persona or {}).get("name") or ""
        try:
            self.persona = json.loads(persona_json) or None
        except Exception:
            self.persona = None
        new = (self.persona or {}).get("name") or ""
        self.config["persona"] = json.dumps(self.persona, ensure_ascii=False) if self.persona else ""
        self._save_config()
        self.core.set_player_persona(self.persona)
        # 玩家卡改名：头像文件跟着改，保持「聊天显示名 = 头像文件名」一致
        eff = new or "你"
        if old and old != new and old != eff:
            avdir = os.path.join(self.save_dir, "avatars")
            for src in (old, "你"):
                for ext in (".png", ".jpg", ".jpeg", ".webp"):
                    f = os.path.join(avdir, src + ext)
                    if os.path.exists(f):
                        try:
                            dst = os.path.join(avdir, eff + ext)
                            if os.path.exists(dst):
                                os.remove(dst)
                            os.rename(f, dst)
                        except Exception:
                            pass
                        break
        return {"ok": True}

    @staticmethod
    def _safe_name(name):
        for ch in [chr(92), "/", ":", "*", "?", '"', "<", ">", "|"]:
            name = name.replace(ch, "_")
        return (name or "").strip()[:60]

    def _parse_role_input(self, name, fields_json):
        """解析角色卡输入：结构化字段 dict / JSON / 旧版纯文本设定 → (legacy, fields, unlocked, advanced)"""
        legacy = ""
        fields = {k: "" for k, _ in ROLE_FIELDS}
        raw = _parse_fields(fields_json)
        if isinstance(fields_json, str) and not raw and fields_json.strip() and not fields_json.strip().startswith("{"):
            legacy = fields_json.strip()
        else:
            legacy = (raw.get("legacy") or raw.get("system_prompt") or "")
            if isinstance(legacy, list):
                legacy = "、".join(str(x) for x in legacy)
            legacy = str(legacy).strip()
            fields = {}
            for k, _ in ROLE_FIELDS:
                v = raw.get(k) or ""
                if isinstance(v, list):
                    v = "、".join(str(x) for x in v if str(x).strip())
                fields[k] = str(v).strip()
        unlocked = False
        if isinstance(fields_json, dict):
            unlocked = bool(fields_json.get("unlocked", False))
        elif raw:
            unlocked = bool(raw.get("unlocked", False))
        # 高级设置（开发者模式）：game / extra_prompt / dev_notes / card_quick_replies
        advanced = None
        adv_raw = None
        if isinstance(fields_json, dict):
            adv_raw = fields_json.get("advanced")
        elif raw:
            adv_raw = raw.get("advanced")
        if isinstance(adv_raw, dict):
            advanced = adv_raw
        elif isinstance(adv_raw, str) and adv_raw.strip():
            try:
                parsed = json.loads(adv_raw)
                if isinstance(parsed, dict):
                    advanced = parsed
            except Exception:
                advanced = None
        return legacy, fields, unlocked, advanced

    @staticmethod
    def _clean_advanced(advanced):
        """规整高级设置字段（只保留已知键，类型安全）"""
        if not isinstance(advanced, dict):
            return None
        game = advanced.get("game")
        if isinstance(game, dict):
            game = {
                "name": str(game.get("name") or "").strip()[:60],
                "rules": str(game.get("rules") or "").strip(),
                "state": str(game.get("state") or "").strip(),
            }
        adv = {
            "game": game if isinstance(game, dict) and (game["rules"] or game["name"]) else None,
            "extra_prompt": str(advanced.get("extra_prompt") or "").strip(),
            "dev_notes": str(advanced.get("dev_notes") or "").strip(),
        }
        mech = advanced.get("mechanics")
        if isinstance(mech, dict):
            cleaned_mech = html_app_clean_mechanics(mech)
            if cleaned_mech:
                adv["mechanics"] = cleaned_mech
        battle = advanced.get("battle")
        if isinstance(battle, dict):
            cleaned_battle = html_app_clean_battle(battle)
            if cleaned_battle:
                adv["battle"] = cleaned_battle
        regex_rules = advanced.get("regex_rules")
        if isinstance(regex_rules, list):
            cleaned_rr = []
            for x in regex_rules:
                if not isinstance(x, dict) or not str(x.get("pattern") or "").strip():
                    continue
                scope = str(x.get("scope") or "both")
                if scope not in ("ai", "user", "both"):
                    scope = "both"
                cleaned_rr.append({
                    "id": str(x.get("id") or "")[:30],
                    "name": str(x.get("name") or x.get("id") or "")[:30],
                    "pattern": str(x.get("pattern") or "").strip(),
                    "replace": str(x.get("replace") or ""),
                    "scope": scope,
                    "enabled": bool(x.get("enabled", True)),
                })
            if cleaned_rr:
                adv["regex_rules"] = cleaned_rr
        qrs = advanced.get("card_quick_replies")
        if isinstance(qrs, list):
            cleaned = []
            for x in qrs:
                if isinstance(x, dict) and x.get("label"):
                    cleaned.append({"label": str(x.get("label")).strip()[:30],
                                    "text": str(x.get("text") or "").strip()})
            adv["card_quick_replies"] = cleaned
        else:
            adv["card_quick_replies"] = []
        if not any(v for k, v in adv.items() if k != "card_quick_replies") and not adv["card_quick_replies"]:
            return None
        return adv

    def api_create_role(self, name, fields_json=None):
        name = (name or "").strip()
        if not name:
            return {"ok": False, "err": "empty"}
        legacy, fields, unlocked, advanced = self._parse_role_input(name, fields_json)
        advanced = self._clean_advanced(advanced)
        prompt = assemble_role_prompt(name, fields, legacy)
        if not prompt:
            return {"ok": False, "err": "empty"}
        fn = self._safe_name(name) + ".json"
        data = {"name": name, "system_prompt": prompt, "legacy": legacy,
                "unlocked": bool(unlocked)}
        for k, v in fields.items():
            if v:
                data[k] = v
        if advanced:
            data["advanced"] = advanced
        save_guard.atomic_write_json(os.path.join(self.save_dir, fn), data)
        self.roles.append({"name": name, "file": fn, "prompt": prompt, "data": data,
                           "fields": fields, "legacy": legacy, "unlocked": bool(unlocked)})
        return {"ok": True}

    def api_update_role(self, name, fields_json=None):
        """编辑人物卡（保持文件名/头像/聊天树/卡数据，只更新设定与结构化字段）"""
        legacy, fields, unlocked, advanced = self._parse_role_input(name, fields_json)
        advanced = self._clean_advanced(advanced)
        prompt = assemble_role_prompt(name, fields, legacy)
        if not prompt:
            return {"ok": False, "err": "empty"}
        for r in self.roles:
            if r["name"] == name:
                r["prompt"] = prompt
                r["fields"] = fields
                r["legacy"] = legacy
                r["unlocked"] = bool(unlocked)
                data = dict(r.get("data") or {})
                data["name"] = name
                data["system_prompt"] = prompt
                data["legacy"] = legacy
                data["unlocked"] = bool(unlocked)
                for k, v in fields.items():
                    if v:
                        data[k] = v
                    else:
                        data.pop(k, None)
                if advanced:
                    data["advanced"] = advanced
                else:
                    data.pop("advanced", None)
                r["data"] = data
                try:
                    save_guard.atomic_write_json(os.path.join(self.save_dir, r["file"]), data)
                except Exception:
                    pass
                if name in self.selected_roles:
                    self._activate_core(reload_tree=False)
                return {"ok": True}
        return {"ok": False, "err": "not found"}

    def _parse_world_params(self, params_json):
        params = {}
        raw = _parse_fields(params_json)
        for key, _label in WORLD_PARAMS:
            v = str(raw.get(key) or "").strip()
            if v:
                params[key] = v
        return params

    def api_create_world(self, name, description, rules_text, entries_json=None, params_json=None):
        """创建世界卡：rules 每行一条；entries 世界书条目；params 世界参数（科技/超自然/物理等）"""
        name = (name or "").strip()
        if not name or not (description or "").strip():
            return {"ok": False, "err": "empty"}
        rules = [x.strip() for x in (rules_text or "").splitlines() if x.strip()]
        data = {"name": name, "description": (description or "").strip(),
                "rules": rules, "entries": self._norm_entries(entries_json),
                "params": self._parse_world_params(params_json)}
        fn = self._safe_name(name) + ".json"
        save_guard.atomic_write_json(os.path.join(self.world_dir, fn), data)
        self.worlds.append(data)
        if name in self.selected_worlds:
            self.core.set_worlds(self._worlds_for_core())
        return {"ok": True}

    def api_update_world(self, name, description, rules_text, entries_json=None, params_json=None):
        """编辑世界卡（按名定位；entries_json/params_json 为 None 时保留原值）"""
        for w in self.worlds:
            if w.get("name") == name:
                w["description"] = (description or "").strip()
                w["rules"] = [x.strip() for x in (rules_text or "").splitlines() if x.strip()]
                if entries_json is not None:
                    w["entries"] = self._norm_entries(entries_json)
                if params_json is not None:
                    w["params"] = self._parse_world_params(params_json)
                try:
                    fn = self._safe_name(name) + ".json"
                    save_guard.atomic_write_json(os.path.join(self.world_dir, fn), w)
                except Exception:
                    pass
                if name in self.selected_worlds:
                    self.core.set_worlds(self._worlds_for_core())
                return {"ok": True}
        return {"ok": False, "err": "not found"}

    def api_delete_world(self, name):
        for w in self.worlds:
            if w.get("name") == name:
                try:
                    os.remove(os.path.join(self.world_dir, self._safe_name(name) + ".json"))
                except Exception:
                    pass
        self.worlds = [w for w in self.worlds if w.get("name") != name]
        self.selected_worlds = [n for n in self.selected_worlds if n != name]
        return {"ok": True}

    def api_ai_draft(self, kind, idea):
        """AI 起草：人物卡/世界卡（需已配置 API Key）"""
        idea = (idea or "").strip()
        if not idea:
            return {"ok": False, "err": "empty idea"}
        client = getattr(self.core, "client", None)
        if not client:
            return {"ok": False, "err": "未配置 API Key，请先在设置里填写"}
        fence = chr(96) * 3
        NL = chr(10)
        if kind == "world":
            prompt = ("根据下面的灵感，设计一个世界卡。只输出 JSON（不要多余解释）：" + NL +
                      '{"name": "世界名", "description": "背景描述", "params": {"tech_level": "科技水平", "supernatural": "超自然体系（无则写：无）", "physics": "物理法则（如：与地球相同）", "time_flow": "时间流速", "climate": "气候环境", "geography": "地理格局", "politics": "政治格局", "economy": "经济体系"}, "rules": ["规则1", "规则2", "规则3"]}' + NL +
                      "灵感：" + idea)
        else:
            prompt = ("根据下面的灵感，撰写一张角色卡。只输出 JSON（不要多余解释）：" + NL +
                      '{"name": "角色名", "appearance": "外貌", "personality": "性格（详写）", "background": "过去经历（详写）", "speech": "说话方式（语气/口癖/句式）", "first_mes": "开场白", "mes_example": "对话示例", "notes": "备注"}' + NL +
                      "灵感：" + idea)
        model = self.config.get("model", "deepseek-v4-flash")
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                stream=False, timeout=90,
            )
            text = (resp.choices[0].message.content or "").strip()
            if fence in text:
                parts = text.split(fence)
                for p in parts:
                    p = p.strip()
                    if p.startswith("json"):
                        p = p[4:].strip()
                    if p.startswith("{"):
                        text = p
                        break
            start = text.find("{")
            end = text.rfind("}")
            if start >= 0 and end > start:
                text = text[start:end + 1]
            data = json.loads(text)
            return {"ok": True, "data": data}
        except Exception as e:
            return {"ok": False, "err": "AI 起草失败：" + str(e)[:200]}

    def api_delete_role(self, name):
        was_selected = name in self.selected_roles
        for r in self.roles:
            if r["name"] == name:
                try:
                    os.remove(os.path.join(self.save_dir, r["file"]))
                except Exception:
                    pass
        self.roles = [r for r in self.roles if r["name"] != name]
        self.selected_roles = [n for n in self.selected_roles if n != name]
        self.config["selected_roles"] = list(self.selected_roles)
        self.config["last_role"] = self.selected_roles[0] if self.selected_roles else ""
        self._save_config()
        if was_selected:
            self._activate_core(reload_tree=True)
        return {"ok": True}

    def api_get_role(self, name):
        """编辑预填：返回人物卡完整数据（结构化字段 + 旧版完整设定）
        旧卡（无结构化字段）→ legacy 预填其原始设定，编辑不丢内容"""
        for r in self.roles:
            if r["name"] == name:
                fields = r.get("fields") or {}
                has_fields = any(v for v in fields.values() if v)
                legacy = r.get("legacy") or ""
                if not legacy and not has_fields:
                    legacy = (r.get("data") or {}).get("system_prompt") or ""
                return {"name": r["name"], "prompt": r["prompt"],
                        "legacy": legacy, "fields": fields,
                        "unlocked": bool(r.get("unlocked", False)),
                        "advanced": (r.get("data") or {}).get("advanced") or None}
        return None

    @staticmethod
    def _entry_hit(entry, text_lower):
        """仅用于前端状态点亮：判断条目是否命中最近一条用户消息"""
        if not entry.get("enabled", True):
            return False
        if entry.get("constant"):
            return True
        kws = entry.get("keywords") or []
        if not kws:
            return False
        mode = (entry.get("match") or "any").lower()
        if mode == "all":
            return all(str(k).lower() in text_lower for k in kws)
        return any(str(k).lower() in text_lower for k in kws)

    def api_get_world(self, name, last_user_text=""):
        """编辑预填：返回世界卡完整数据（附带条目触发状态）"""
        for w in self.worlds:
            if w.get("name") == name:
                w = dict(w)
                txt = (last_user_text or "").lower()
                ent = []
                for e in w.get("entries", []):
                    e = dict(e)
                    e["triggered"] = self._entry_hit(e, txt)
                    ent.append(e)
                w["entries"] = ent
                return w
        return None

    def api_get_avatar(self, name):
        """自定义头像 base64（无则 null，前端画首字圆头像）"""
        avdir = os.path.join(self.save_dir, "avatars")
        for ext in (".png", ".jpg", ".jpeg", ".webp"):
            p = os.path.join(avdir, str(name) + ext)
            if os.path.exists(p):
                try:
                    with open(p, "rb") as f:
                        mime = "image/" + ("jpeg" if ext == ".jpg" else ext[1:])
                        return "data:" + mime + ";base64," + base64.b64encode(f.read()).decode()
                except Exception:
                    return None
        return None

    def api_set_avatar(self, name, b64, ext):
        try:
            raw = base64.b64decode(b64.split(",", 1)[-1])
            avdir = os.path.join(self.save_dir, "avatars")
            os.makedirs(avdir, exist_ok=True)
            ext = (ext or "png").replace("jpeg", "jpg")
            with open(os.path.join(avdir, str(name) + "." + ext), "wb") as f:
                f.write(raw)
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)}

    # ---------- 滑条 / 编辑 / 分支 ----------
    def api_regenerate(self, seq):
        """重新生成：为最后一条 AI 消息生成新候选（滑条）"""
        msg = next((m for m in self.messages if m["seq"] == seq), None)
        if not msg or msg["kind"] != "ai":
            return {"ok": False, "err": "只能对 AI 消息重新生成"}
        if self.busy:
            return {"ok": False, "err": "busy"}
        node = self.core.tree.nodes.get(msg["node_id"])
        if not node or not node.parent_id or node.parent_id not in self.core.tree.nodes:
            return {"ok": False, "err": "节点不存在"}
        self._start_fetch(node.parent_id)
        return {"ok": True}

    def api_switch_swipe(self, seq, index):
        """切换滑条：选择父用户节点下的第 index 条 assistant 候选"""
        msg = next((m for m in self.messages if m["seq"] == seq), None)
        if not msg or msg["kind"] != "ai":
            return {"ok": False, "err": "无效消息"}
        node = self.core.tree.nodes.get(msg["node_id"])
        if not node or not node.parent_id or node.parent_id not in self.core.tree.nodes:
            return {"ok": False, "err": "节点不存在"}
        sibs = [cid for cid in self.core.tree.nodes[node.parent_id].children_ids
                if self.core.tree.nodes.get(cid) and self.core.tree.nodes[cid].role == "assistant"]
        try:
            target = sibs[int(index)]
        except (ValueError, IndexError):
            return {"ok": False, "err": "滑条越界"}
        self.core.tree.current_leaf_id = target
        self._save_tree()
        self._rebuild_messages()
        return {"ok": True}

    def api_edit_message(self, seq, new_content):
        """编辑消息：AI 消息原地改；用户消息开新分支并自动重新生成"""
        new_content = (new_content or "").strip()
        if not new_content:
            return {"ok": False, "err": "内容不能为空"}
        msg = next((m for m in self.messages if m["seq"] == seq), None)
        if not msg:
            return {"ok": False, "err": "无效消息"}
        node = self.core.tree.nodes.get(msg["node_id"])
        if not node:
            return {"ok": False, "err": "节点不存在"}
        if node.role == "assistant":
            node.content = new_content
            self._save_tree()
            self._rebuild_messages()
            return {"ok": True}
        if self.busy:
            return {"ok": False, "err": "busy"}
        new_id = self.core.tree.add_node("user", new_content,
                                         parent_id=node.parent_id,
                                         metadata=dict(node.metadata or {}))
        if node.id in self.node_images:
            self.node_images[new_id] = self.node_images[node.id]
        self._start_fetch(new_id)
        return {"ok": True}

    def api_branches(self):
        """列出可切换的其他分支（叶子节点）"""
        tree = self.core.tree
        out = []
        for nid, node in tree.nodes.items():
            if node.role == "system" or node.children_ids:
                continue
            if nid == tree.current_leaf_id:
                continue
            previews = []
            n = node
            depth = 0
            while n is not None and len(previews) < 2:
                if n.role != "system":
                    previews.append(n.content)
                depth += 1
                if n.parent_id and n.parent_id in tree.nodes:
                    n = tree.nodes[n.parent_id]
                else:
                    break
            previews.reverse()
            out.append({"node_id": nid, "kind": node.role, "depth": depth,
                        "preview": " / ".join((p or "").replace(chr(10), " ")[:40] for p in previews)})
        out.sort(key=lambda x: -x["depth"])
        return out

    def api_switch_branch(self, node_id):
        tree = self.core.tree
        if not node_id or node_id not in tree.nodes:
            return {"ok": False, "err": "节点不存在"}
        tree.current_leaf_id = node_id
        self.core.restore_mechanisms(node_id)
        self._save_tree()
        self._rebuild_messages()
        return {"ok": True}

    def api_tree(self):
        """树状回溯：返回完整历史树（主线平铺 + 分支收纳）。
        当前路径（主线）上的节点 on_path=True（不缩进、不右窜）；
        离线分支以 branch_root 标识：branch_root==自身 的行是「分支收纳行」，
        其 branch_size 为分支节点数，展开后成员按 branch_depth（分支内深度）显示。"""
        tree = self.core.tree
        nodes = tree.nodes
        # 当前路径（root → 当前叶子），主线 = 路径上的节点
        path = set()
        nid = tree.current_leaf_id
        while nid and nid in nodes:
            path.add(nid)
            nid = nodes[nid].parent_id

        def subtree_size(nid):
            node = nodes.get(nid)
            if not node:
                return 0
            return 1 + sum(subtree_size(c) for c in node.children_ids)

        out = []
        branch_sizes = {}

        def walk(nid, in_branch, branch_depth):
            node = nodes.get(nid)
            if not node:
                return
            on_path = nid in path
            if on_path:
                in_branch = None
                branch_depth = 0
            elif in_branch is None:
                in_branch = nid
                branch_depth = 0
                branch_sizes[nid] = subtree_size(nid)
            else:
                branch_depth += 1
            out.append({
                "id": nid,
                "role": node.role,
                "content": (node.content or "").replace("\n", " ")[:60],
                "on_path": on_path,
                "branch_root": in_branch,
                "branch_depth": branch_depth,
                "branch_size": branch_sizes.get(in_branch, 0),
                "is_current": nid == tree.current_leaf_id,
                "is_leaf": not node.children_ids,
            })
            for c in node.children_ids:
                walk(c, in_branch, branch_depth)

        if tree.root_id:
            walk(tree.root_id, None, 0)
        return out

    def api_backtrack(self, node_id):
        """树状回溯：把对话定位到任意节点（从这里继续/切换分支）"""
        tree = self.core.tree
        if not node_id or node_id not in tree.nodes:
            return {"ok": False, "err": "节点不存在"}
        if node_id == tree.current_leaf_id:
            return {"ok": True}
        tree.current_leaf_id = node_id
        self.core.restore_mechanisms(node_id)
        self._save_tree()
        self._rebuild_messages()
        return {"ok": True}

    # ---------- 酒馆角色卡导入导出 ----------
    def _do_import_card(self, path):
        try:
            with open(path, "rb") as f:
                raw = f.read()
        except Exception as e:
            return {"ok": False, "err": "读取失败：" + str(e)}
        low = path.lower()
        if low.endswith(".png") or low.endswith(".webp"):
            data = card_compat.png_extract_card(raw)
            if data is None:
                return {"ok": False, "err": "这张图片里没有找到角色卡（chara/ccv3 块）"}
        else:
            try:
                data = json.loads(raw.decode("utf-8-sig"))
            except Exception as e:
                return {"ok": False, "err": "JSON 解析失败：" + str(e)}
        conv = card_compat.to_dick(data)
        if not conv:
            return {"ok": False, "err": "无法识别的角色卡格式（需要 v1/v2/v3 或 DICK 格式）"}
        name = conv["name"]
        base_name = name
        i = 2
        while any(r["name"] == name for r in self.roles):
            name = base_name + "_" + str(i)
            i += 1
        fn = self._safe_name(name) + ".json"
        saved = {"name": name, "system_prompt": conv["system_prompt"]}
        if isinstance(conv.get("card_data"), dict):
            saved["card_data"] = conv["card_data"]
        # 完全适配：结构化字段拆分（可编辑）+ 备用开场白
        if isinstance(conv.get("fields"), dict) and conv["fields"]:
            for k, v in conv["fields"].items():
                saved[k] = v
        if conv.get("alternate_greetings"):
            saved["alternate_greetings"] = conv["alternate_greetings"]
        save_guard.atomic_write_json(os.path.join(self.save_dir, fn), saved)
        self.roles.append({"name": name, "file": fn, "prompt": conv["system_prompt"], "data": saved})
        if low.endswith((".png", ".webp")):
            try:
                avdir = os.path.join(self.save_dir, "avatars")
                os.makedirs(avdir, exist_ok=True)
                ext = "webp" if low.endswith(".webp") else "png"
                with open(os.path.join(avdir, self._safe_name(name) + "." + ext), "wb") as f:
                    f.write(raw)
            except Exception:
                pass
        # 完全适配：酒馆 v2 内嵌世界书（extensions.world）→ DICK 世界卡
        world_note = ""
        entries = conv.get("world_entries") or []
        if entries:
            try:
                wn = name + " 的世界书"
                found = next((w for w in self.worlds if w.get("name") == wn), None)
                if found:
                    merged = found.get("entries") or []
                    existing_ids = {str(e.get("id")) for e in merged}
                    for e in entries:
                        if str(e.get("id")) not in existing_ids:
                            merged.append(e)
                    found["entries"] = merged
                    save_guard.atomic_write_json(
                        os.path.join(self.world_dir, self._safe_name(wn) + ".json"), found)
                else:
                    wdata = {"name": wn,
                             "description": "从角色卡「" + name + "」导入的酒馆世界书",
                             "rules": [], "entries": entries, "params": {}}
                    save_guard.atomic_write_json(
                        os.path.join(self.world_dir, self._safe_name(wn) + ".json"), wdata)
                    self.worlds.append(wdata)
                world_note = "，世界书 " + str(len(entries)) + " 条 → 世界卡「" + wn + "」"
            except Exception as e:
                print(f"[导入] 世界书写入失败: {e}")
        return {"ok": True, "name": name, "note": world_note}

    def api_import_card(self):
        """文件对话框导入酒馆角色卡（.png/.webp/.json）"""
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            res = win.create_file_dialog(
                webview.FileDialog.OPEN, allow_multiple=False,
                file_types=("角色卡 (*.png;*.webp;*.json)", "All files (*.*)"))
        except Exception as e:
            return {"ok": False, "err": "对话框失败：" + str(e)}
        if not res:
            return {"ok": False, "err": "cancelled"}
        return self._do_import_card(res[0])

    def _do_export_card(self, name, fmt, path):
        role = next((r for r in self.roles if r["name"] == name), None)
        if not role:
            return {"ok": False, "err": "角色不存在"}
        # 关联世界卡（`<角色名> 的世界书`）→ 导出时写回酒馆 extensions.world（无损反向）
        world_entries = []
        try:
            wn = name + " 的世界书"
            w = next((x for x in self.worlds if x.get("name") == wn), None)
            if w:
                world_entries = w.get("entries") or []
        except Exception:
            pass
        card = card_compat.dick_to_v2(name, role["prompt"],
                                      (role.get("data") or {}).get("card_data"),
                                      world_entries=world_entries)
        try:
            if fmt == "json":
                save_guard.atomic_write_json(path, card)
            else:
                png = None
                avdir = os.path.join(self.save_dir, "avatars")
                for ext in (".png", ".jpg", ".jpeg", ".webp"):
                    p = os.path.join(avdir, self._safe_name(name) + ext)
                    if os.path.exists(p):
                        with open(p, "rb") as f:
                            png = f.read()
                        break
                if png is None:
                    png = card_compat.placeholder_png(name)
                with open(path, "wb") as f:
                    f.write(card_compat.png_embed_card(png, card))
        except Exception as e:
            return {"ok": False, "err": "导出失败：" + str(e)}
        return {"ok": True}

    def api_export_card(self, name, fmt):
        """导出角色卡：fmt = json | png（酒馆 v2 格式，PNG 嵌卡）"""
        fmt = (fmt or "json").lower()
        if fmt not in ("json", "png"):
            return {"ok": False, "err": "格式必须是 json 或 png"}
        ext = "json" if fmt == "json" else "png"
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            res = win.create_file_dialog(
                webview.FileDialog.SAVE,
                save_filename=self._safe_name(name) + "." + ext,
                file_types=("JSON (*.json)",) if fmt == "json" else ("PNG (*.png)",))
        except Exception as e:
            return {"ok": False, "err": "对话框失败：" + str(e)}
        if not res:
            return {"ok": False, "err": "cancelled"}
        path = res[0]
        if fmt == "json" and not path.lower().endswith(".json"):
            path += ".json"
        if fmt == "png" and not path.lower().endswith(".png"):
            path += ".png"
        return self._do_export_card(name, fmt, path)

    # ---------- 世界书条目 ----------
    def _norm_entries(self, entries_json):
        """规范化世界书条目列表（字段齐全、类型安全）"""
        try:
            raw = json.loads(entries_json) if isinstance(entries_json, str) else (entries_json or [])
        except Exception:
            raw = []
        if not isinstance(raw, list):
            raw = []

        def num(v, d):
            try:
                return float(v)
            except (TypeError, ValueError):
                return d

        entries = []
        for e in raw:
            if not isinstance(e, dict):
                continue
            keywords = e.get("keywords", [])
            if isinstance(keywords, str):
                keywords = [k.strip() for k in keywords.replace("，", ",").split(",") if k.strip()]
            keywords = [str(k).strip() for k in keywords if str(k).strip()]
            content = str(e.get("content", "")).strip()
            if not content:
                continue
            entries.append({
                "id": str(e.get("id") or ""),
                "keywords": keywords,
                "content": content,
                "match": str(e.get("match") or "any").lower(),
                "weight": num(e.get("weight"), 100),
                "probability": num(e.get("probability"), 100),
                "depth": int(num(e.get("depth"), 1)),
                "enabled": bool(e.get("enabled", True)),
                "constant": bool(e.get("constant", False)),
            })
        return entries

# ---------- 创意工坊（联网版，端口自 Tk 版） ----------
    def _ws_config_file(self):
        return os.path.join(BASE_DIR, "workshop_config.json")

    def _ws_load_config(self):
        try:
            with open(self._ws_config_file(), "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}

    def _ws_save_config(self, server_url, api_key, proxy=None):
        with open(self._ws_config_file(), "w", encoding="utf-8") as f:
            json.dump({"server_url": (server_url or "").strip(),
                       "api_key": (api_key or "").strip(),
                       "proxy": (proxy or "").strip()}, f, ensure_ascii=False, indent=2)

    def _ws_headers(self):
        key = (self._ws_load_config().get("api_key") or "").strip()
        return {"X-API-Key": key} if key else {}

    def _ws_proxies(self):
        """从配置读取代理（支持 http/https/socks5），返回 requests 代理字典或 None"""
        proxy = (self._ws_load_config().get("proxy") or "").strip()
        if not proxy:
            return None
        if "://" not in proxy:
            proxy = "http://" + proxy
        return {"http": proxy, "https": proxy}

    def _ws_url(self, path):
        base = self._ws_active_server()
        return base + path

    def _ws_active_server(self):
        """自动部署：返回当前可用的工坊服务器地址。
        优先使用配置的 server_url；未配置或不可用时自动探测
        ① 本机服务器 ② 默认隧道地址。结果缓存 30 秒避免每次请求都探测。"""
        import time as _t
        cfg_url = (self._ws_load_config().get("server_url") or "").strip().rstrip("/")
        now = _t.time()
        # 缓存有效期内直接返回
        cached = getattr(self, "_ws_active_cache", None)
        if cached and now - cached[1] < 30:
            return cached[0]
        candidates = []
        # 本机服务器优先（最快最稳）
        candidates.append("http://127.0.0.1:5000")
        # 默认隧道地址（保活脚本维护的稳定入口，可能变化但当前可用）
        candidates.append("https://referrals-lambda-geographical-says.trycloudflare.com")
        if cfg_url:
            candidates.append(cfg_url)
        import requests as _req
        for base in candidates:
            try:
                r = _req.get(base + "/api/health", timeout=2.5,
                             headers=self._ws_headers(), proxies=self._ws_proxies())
                if r.status_code < 400:
                    self._ws_active_cache = (base, now)
                    return base
            except Exception:
                continue
        # 全部失败：退回配置地址（后续请求会报错并提示）
        self._ws_active_cache = (cfg_url or candidates[1], now)
        return cfg_url or candidates[1]

    def _reload_roles(self):
        """从磁盘重建角色列表（保留选中项）"""
        self.roles = []
        try:
            for fn in sorted(os.listdir(self.save_dir)):
                if not fn.endswith(".json"):
                    continue
                try:
                    with open(os.path.join(self.save_dir, fn), "r", encoding="utf-8") as f:
                        data = json.load(f)
                    if isinstance(data, dict) and data.get("system_prompt"):
                        fields = {k: (data.get(k) or "") for k, _ in ROLE_FIELDS}
                        legacy = data.get("legacy") or ""
                        prompt = assemble_role_prompt(data.get("name") or fn[:-5], fields, legacy) \
                            or data.get("system_prompt", "")
                        self.roles.append({"name": data.get("name") or fn[:-5], "file": fn,
                                           "prompt": prompt, "data": data,
                                           "fields": fields, "legacy": legacy})
                except Exception:
                    continue
        except Exception:
            pass
        self.selected_roles = [n for n in self.selected_roles if any(r["name"] == n for r in self.roles)]

    def _reload_worlds(self):
        """从磁盘重建世界列表（保留选中项）"""
        self.worlds = []
        try:
            for fn in sorted(os.listdir(self.world_dir)):
                if not fn.endswith(".json"):
                    continue
                try:
                    with open(os.path.join(self.world_dir, fn), "r", encoding="utf-8") as f:
                        data = json.load(f)
                    if isinstance(data, dict) and data.get("name"):
                        self.worlds.append(data)
                except Exception:
                    continue
        except Exception:
            pass
        self.selected_worlds = [n for n in self.selected_worlds if any(w.get("name") == n for w in self.worlds)]

    def api_workshop_state(self):
        cfg = self._ws_load_config()
        roles = [fn for fn in sorted(os.listdir(self.save_dir)) if fn.endswith(".json")]
        worlds = [fn for fn in sorted(os.listdir(self.world_dir)) if fn.endswith(".json")]
        active = ""
        try:
            active = self._ws_active_server()
        except Exception:
            pass
        return {"server_url": cfg.get("server_url", ""), "active_server": active,
                "has_key": bool(cfg.get("api_key")),
                "proxy": cfg.get("proxy", ""),
                "local_roles": roles, "local_worlds": worlds}

    def api_workshop_save_conn(self, server_url, api_key, proxy=None):
        self._ws_save_config(server_url or "", api_key or "", proxy or "")
        return {"ok": True}

    def api_workshop_detect_proxy(self):
        """自动检测本机代理：1) Windows 系统代理设置 2) 常见代理软件端口"""
        import socket
        found = []
        # 1) Windows 系统代理（IE/系统设置里开的代理）
        try:
            import winreg
            key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                                 r"SoftwareMicrosoftWindowsCurrentVersionInternet Settings")
            enable, _ = winreg.QueryValueEx(key, "ProxyEnable")
            server, _ = winreg.QueryValueEx(key, "ProxyServer")
            winreg.CloseKey(key)
            if enable and server:
                srv = str(server).strip()
                # 可能形如 "127.0.0.1:7890" 或 "http=127.0.0.1:7890;https=..."; 取第一个
                if "=" in srv:
                    srv = srv.split(";")[0].split("=", 1)[1].strip()
                if srv and "://" not in srv:
                    srv = "http://" + srv
                found.append(("系统代理", srv))
        except Exception:
            pass
        # 2) 常见代理软件本地端口探测
        common_ports = [7890, 7891, 10809, 10808, 1080, 8888, 1087, 8080, 20171, 9910]
        for port in common_ports:
            try:
                s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                s.settimeout(0.25)
                if s.connect_ex(("127.0.0.1", port)) == 0:
                    found.append((f"本地端口 :{port}", f"http://127.0.0.1:{port}"))
                s.close()
            except Exception:
                pass
        # 去重（同一地址只保留一次）
        seen = set()
        result = []
        for label, url in found:
            if url not in seen:
                seen.add(url)
                result.append({"label": label, "url": url})
        return {"ok": True, "found": result}

    def api_workshop_test(self):
        import requests
        try:
            r = requests.get(self._ws_url("/api/health"), timeout=6, headers=self._ws_headers(), proxies=self._ws_proxies())
            if r.status_code >= 400:
                return {"ok": False, "err": "HTTP " + str(r.status_code)}
            try:
                health = r.json() if isinstance(r.json(), dict) else {}
            except Exception:
                health = {}
            stats = {}
            try:
                sr = requests.get(self._ws_url("/api/stats"), timeout=6, headers=self._ws_headers(), proxies=self._ws_proxies())
                if sr.status_code < 400:
                    stats = sr.json()
            except Exception:
                pass
            return {"ok": True, "health": health, "stats": stats}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_workshop_list(self):
        import requests
        try:
            cr = requests.get(self._ws_url("/api/cards/list"), timeout=8, headers=self._ws_headers(), proxies=self._ws_proxies())
            wr = requests.get(self._ws_url("/api/worlds/list"), timeout=8, headers=self._ws_headers(), proxies=self._ws_proxies())
            if cr.status_code >= 400 or wr.status_code >= 400:
                return {"ok": False, "err": "HTTP " + str(cr.status_code) + " / " + str(wr.status_code)}
            return {"ok": True, "cards": cr.json(), "worlds": wr.json()}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_workshop_search(self, query, rtype):
        import requests
        try:
            params = {}
            if query:
                params["q"] = query
            if rtype:
                params["type"] = rtype
            r = requests.get(self._ws_url("/api/search"), params=params, timeout=8, headers=self._ws_headers(), proxies=self._ws_proxies())
            if r.status_code >= 400:
                return {"ok": False, "err": "HTTP " + str(r.status_code)}
            return {"ok": True, "results": r.json()}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_workshop_download(self, res_id, rtype, filename):
        import requests
        try:
            kind = "cards" if rtype == "角色卡" else "worlds"
            target_dir = self.save_dir if rtype == "角色卡" else self.world_dir
            r = requests.get(self._ws_url("/api/" + kind + "/" + str(res_id)), timeout=30,
                             headers=self._ws_headers(), proxies=self._ws_proxies())
            if r.status_code >= 400:
                return {"ok": False, "err": "HTTP " + str(r.status_code)}
            fname = (filename or (str(res_id) + ".json")).replace("..", "_")
            if not fname.endswith(".json"):
                fname += ".json"
            dst = os.path.join(target_dir, fname)
            base, ext = os.path.splitext(dst)
            i = 1
            while os.path.exists(dst):
                dst = base + "_" + str(i) + ext
                i += 1
            with open(dst, "wb") as f:
                f.write(r.content)
            if rtype == "角色卡":
                self._reload_roles()
            else:
                self._reload_worlds()
            return {"ok": True, "file": os.path.basename(dst)}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_workshop_like(self, res_id, rtype):
        import requests
        try:
            kind = "cards" if rtype == "角色卡" else "worlds"
            r = requests.post(self._ws_url("/api/" + kind + "/" + str(res_id) + "/like"),
                              timeout=8, headers=self._ws_headers(), proxies=self._ws_proxies())
            if r.status_code >= 400:
                return {"ok": False, "err": "HTTP " + str(r.status_code)}
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_workshop_delete(self, res_id, rtype):
        import requests
        try:
            kind = "cards" if rtype == "角色卡" else "worlds"
            r = requests.delete(self._ws_url("/api/" + kind + "/" + str(res_id)),
                                timeout=8, headers=self._ws_headers(), proxies=self._ws_proxies())
            if r.status_code >= 400:
                return {"ok": False, "err": "HTTP " + str(r.status_code)}
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_workshop_upload(self, rtype, name):
        import requests
        try:
            target_dir = self.save_dir if rtype == "角色卡" else self.world_dir
            fname = self._safe_name(name) + ".json"
            path = os.path.join(target_dir, fname)
            if not os.path.isfile(path):
                return {"ok": False, "err": "本地文件不存在：" + fname}
            with open(path, "rb") as f:
                r = requests.post(
                    self._ws_url("/api/cards/upload" if rtype == "角色卡" else "/api/worlds/upload"),
                    files={"file": (fname, f, "application/json")}, data={"name": name},
                    timeout=30, headers=self._ws_headers(), proxies=self._ws_proxies())
            if r.status_code >= 400:
                return {"ok": False, "err": "HTTP " + str(r.status_code)}
            return {"ok": True, "message": r.text[:200]}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_workshop_preview(self, rtype, filename):
        target_dir = self.save_dir if rtype == "角色卡" else self.world_dir
        filename = os.path.basename((filename or "").replace("..", "_"))
        path = os.path.join(target_dir, filename)
        if not os.path.isfile(path):
            return {"ok": False, "err": "文件不存在"}
        try:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read(8000)
            try:
                json.loads(text)
            except Exception:
                pass
            return {"ok": True, "preview": text}
        except Exception as e:
            return {"ok": False, "err": str(e)}

    def api_workshop_export(self, rtype, filename):
        """本地文件导出到用户选择的位置"""
        target_dir = self.save_dir if rtype == "角色卡" else self.world_dir
        filename = os.path.basename((filename or "").replace("..", "_"))
        src = os.path.join(target_dir, filename)
        if not os.path.isfile(src):
            return {"ok": False, "err": "文件不存在"}
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            res = win.create_file_dialog(webview.FileDialog.SAVE, save_filename=filename,
                                         file_types=("JSON (*.json)",))
            if not res:
                return {"ok": False, "err": "cancelled"}
            import shutil
            shutil.copy2(src, res[0])
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_workshop_plugins(self):
        """拉取工坊插件列表 + 本地已安装插件名（标记用）"""
        import requests
        try:
            r = requests.get(self._ws_url("/api/plugins/list"), timeout=8,
                             headers=self._ws_headers(), proxies=self._ws_proxies())
            if r.status_code >= 400:
                return {"ok": False, "err": "HTTP " + str(r.status_code)}
            plugins = r.json()
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}
        local = set()
        for d in getattr(self.plugin_manager, "plugin_dirs", []) or []:
            if os.path.isdir(d):
                for fn in os.listdir(d):
                    if fn.endswith(".py") and not fn.startswith("_"):
                        local.add(fn[:-3])
        return {"ok": True, "plugins": plugins, "local": sorted(local)}

    def api_workshop_install_plugin(self, plugin_id):
        """下载插件到 plugins/ 目录并热加载"""
        import requests
        try:
            r = requests.get(self._ws_url("/api/plugins/" + str(plugin_id)), timeout=30,
                             headers=self._ws_headers(), proxies=self._ws_proxies())
            if r.status_code >= 400:
                return {"ok": False, "err": "HTTP " + str(r.status_code)}
            info = None
            try:
                lr = requests.get(self._ws_url("/api/plugins/list"), timeout=8,
                                  headers=self._ws_headers(), proxies=self._ws_proxies())
                if lr.status_code < 400:
                    info = next((p for p in lr.json() if str(p.get("id")) == str(plugin_id)), None)
            except Exception:
                pass
            fname = (info or {}).get("original_name") or (str(plugin_id) + ".py")
            fname = os.path.basename(str(fname).replace("..", "_"))
            if not fname.endswith(".py"):
                fname += ".py"
            target_dir = getattr(self.plugin_manager, "plugin_dir", None) or "plugins"
            os.makedirs(target_dir, exist_ok=True)
            dst = os.path.join(target_dir, fname)
            with open(dst, "wb") as f:
                f.write(r.content)
            try:
                self.plugin_manager.reload_plugins()
            except Exception as e:
                print(f"[Plugin] 重载失败: {e}")
            return {"ok": True, "name": (info or {}).get("name") or fname, "file": fname}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_workshop_delete_local(self, rtype, filename):
        target_dir = self.save_dir if rtype == "角色卡" else self.world_dir
        filename = os.path.basename((filename or "").replace("..", "_"))
        path = os.path.join(target_dir, filename)
        if not os.path.isfile(path):
            return {"ok": False, "err": "文件不存在"}
        try:
            os.remove(path)
            if rtype == "角色卡":
                self._reload_roles()
            else:
                self._reload_worlds()
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_workshop_open_folder(self, which):
        try:
            if which == "roles":
                target = self.save_dir
            elif which == "worlds":
                target = self.world_dir
            elif which == "plugins":
                target = getattr(self.plugin_manager, "plugin_dir", None) or "plugins"
            else:
                target = self.save_dir
            os.startfile(target)
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)}

    def api_open_path(self, path):
        """在资源管理器中打开/定位文件（打包产物定位用）"""
        try:
            if path and os.path.exists(path):
                os.startfile(os.path.dirname(os.path.abspath(path)))
                return {"ok": True}
            return {"ok": False, "err": "not found"}
        except Exception as e:
            return {"ok": False, "err": str(e)}

# ---------- 前版功能补齐：预算 / 自动接话 / 字体 / 插件管理 / 文档 / 导出 ----------
    def api_toggle_unlock(self, name):
        """破甲模式快速开关（角色级：任一启用则对话注入破甲提示）"""
        for r in self.roles:
            if r["name"] == name:
                r["unlocked"] = not r.get("unlocked", False)
                data = dict(r.get("data") or {})
                data["unlocked"] = r["unlocked"]
                r["data"] = data
                try:
                    save_guard.atomic_write_json(os.path.join(self.save_dir, r["file"]), data)
                except Exception:
                    pass
                if name in self.selected_roles:
                    self._activate_core(reload_tree=False)
                return {"ok": True, "unlocked": r["unlocked"]}
        return {"ok": False, "err": "not found"}

    def api_set_budget(self, tokens):
        try:
            tokens = int(tokens or 0)
        except (TypeError, ValueError):
            tokens = 0
        tokens = max(0, tokens)
        self.config["context_budget"] = tokens
        self.core.set_context_budget(tokens)
        self._save_config()
        return {"ok": True, "budget": tokens}

    def api_set_auto_turn(self, on):
        self.auto_turn = bool(on)
        self.config["auto_turn"] = self.auto_turn
        self._save_config()
        return {"ok": True, "auto_turn": self.auto_turn}

    def api_set_font(self, idx):
        try:
            idx = int(idx or 0)
        except (TypeError, ValueError):
            idx = 0
        idx = max(0, min(2, idx))
        self.font_size = idx
        self.config["ui_font"] = idx
        self._save_config()
        return {"ok": True, "font": idx}

    def api_set_plugin(self, name, enabled):
        p = self.plugin_manager.get_plugin(name)
        if not p:
            return {"ok": False, "err": "插件不存在"}
        p.enabled = bool(enabled)
        states = dict(self.config.get("plugin_states") or {})
        states[p.name] = p.enabled
        self.config["plugin_states"] = states
        self._save_config()
        return {"ok": True}

    def api_set_plugin_setting(self, name, key, value):
        p = self.plugin_manager.get_plugin(name)
        if not p:
            return {"ok": False, "err": "插件不存在"}
        schema = {item.get("key"): item for item in (getattr(p, "settings_schema", None) or [])}
        item = schema.get(key)
        if not item:
            return {"ok": False, "err": "未知设置项"}
        vtype = item.get("type", "text")
        try:
            if vtype == "bool":
                value = bool(value)
            elif vtype == "int":
                value = int(float(value or 0))
            else:
                value = str(value or "")
        except (TypeError, ValueError):
            return {"ok": False, "err": "数值格式错误"}
        p.set_setting(key, value)
        return {"ok": True}

    def api_read_document(self, path=None):
        """读入 Word/Excel 文档供 AI 分析（文件对话框或给定路径）"""
        if not path:
            try:
                import webview
                win = webview.windows[0] if getattr(webview, "windows", None) else None
                if win is None:
                    return {"ok": False, "err": "窗口未就绪"}
                res = win.create_file_dialog(
                    webview.FileDialog.OPEN, allow_multiple=False,
                    file_types=("文档 (*.docx;*.xlsx;*.txt)", "All files (*.*)"))
                if not res:
                    return {"ok": False, "err": "cancelled"}
                path = res[0]
            except Exception as e:
                return {"ok": False, "err": "对话框失败：" + str(e)}
        path = path.strip().strip('"')
        try:
            import doc_reader
            text = doc_reader.read_document(path)
        except ValueError as e:
            # 兼容 .txt 直读
            if path.lower().endswith(".txt"):
                try:
                    with open(path, "r", encoding="utf-8", errors="replace") as f:
                        text = f.read()
                except Exception as e2:
                    return {"ok": False, "err": str(e2)}
            else:
                return {"ok": False, "err": str(e)}
        except Exception as e:
            return {"ok": False, "err": "读取失败：" + str(e)[:200]}
        if not text or not text.strip():
            return {"ok": False, "err": "文档内容为空"}
        cap = 20000
        if len(text) > cap:
            text = text[:cap] + chr(10) + "…（文档过长，已截断，可用 /readclear 清除）"
        self.core.set_document_context(text)
        self.loaded_document = {"name": os.path.basename(path), "chars": len(text)}
        self._append_sys("📄 已读入文档：" + self.loaded_document["name"] +
                         "（" + str(len(text)) + " 字符），后续对话自动参考")
        return {"ok": True, "name": self.loaded_document["name"], "chars": len(text)}

    def api_clear_document(self):
        self.core.clear_document_context()
        self.loaded_document = None
        self._append_sys("🧹 已清除文档上下文")
        return {"ok": True}

    def api_export_chat(self):
        """把当前聊天记录导出为 Word 文档"""
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            res = win.create_file_dialog(webview.FileDialog.SAVE,
                                         save_filename="聊天记录.docx",
                                         file_types=("Word (*.docx)",))
            if not res:
                return {"ok": False, "err": "cancelled"}
            path = res[0]
            if not path.lower().endswith(".docx"):
                path += ".docx"
            from docx import Document
            doc = Document()
            doc.add_heading("Direct-Interface Cork-bore Kit · 聊天记录", level=1)
            for m in self.messages:
                who = m.get("speaker") or ("你" if m.get("kind") == "user" else "AI")
                doc.add_paragraph("")
                p = doc.add_paragraph()
                run = p.add_run(who + "：")
                run.bold = True
                doc.add_paragraph(m.get("content") or "")
            doc.save(path)
            return {"ok": True, "file": os.path.basename(path)}
        except Exception as e:
            return {"ok": False, "err": "导出失败：" + str(e)[:200]}

    # ================= CODEX 专属 GALGAME（PC 端独有） =================
    # 傻瓜化：选文件夹自动归类立绘/背景/音乐/配音 → 生成剧本模板 → 播放器直接能播。
    # JSON 功能：codex.json 可视化编辑 + 校验 + zip 导入导出。

    def _codex_pkg_path(self, name):
        return os.path.join(self.codex_dir, codex_core._safe_name(name))

    def api_codex_list(self):
        """列出所有 CODEX 包：name / intro / 场景数 / 资源统计"""
        out = []
        try:
            for fn in sorted(os.listdir(self.codex_dir)):
                d = os.path.join(self.codex_dir, fn)
                if not os.path.isdir(d):
                    continue
                info = {"name": fn, "intro": "", "scenes": 0,
                        "sprites": 0, "bg": 0, "bgm": 0, "voice": 0,
                        "has_script": False}
                jp = os.path.join(d, "codex.json")
                if os.path.isfile(jp):
                    try:
                        with open(jp, "r", encoding="utf-8") as f:
                            data = json.load(f)
                        info["has_script"] = True
                        info["intro"] = str(data.get("intro") or "")[:120]
                        info["scenes"] = len(data.get("scenes") or [])
                    except Exception:
                        pass
                for kind in codex_core.SUBDIRS:
                    kd = os.path.join(d, kind)
                    if os.path.isdir(kd):
                        try:
                            info[kind] = len([x for x in os.listdir(kd)
                                              if os.path.isfile(os.path.join(kd, x))])
                        except Exception:
                            pass
                out.append(info)
        except Exception:
            pass
        return {"ok": True, "packages": out}

    def api_codex_create(self, name):
        """新建 CODEX 包：生成示例剧本 + 空资源目录"""
        name = (name or "").strip()[:60]
        if not name:
            return {"ok": False, "err": "empty"}
        dst = self._codex_pkg_path(name)
        try:
            os.makedirs(dst, exist_ok=True)
            for k in codex_core.SUBDIRS:
                os.makedirs(os.path.join(dst, k), exist_ok=True)
            jp = os.path.join(dst, "codex.json")
            if not os.path.isfile(jp):
                with open(jp, "w", encoding="utf-8") as f:
                    json.dump(codex_core.make_template(name), f,
                              ensure_ascii=False, indent=2)
            return {"ok": True, "name": os.path.basename(dst)}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_codex_sample(self, name="示例·夏日祭"):
        """一键示例包：生成带占位立绘/背景/音乐/配音的真实可播 CODEX 包"""
        ok, dst = codex_core.make_sample_package(self.codex_dir, name)
        if not ok:
            return {"ok": False, "err": "生成失败：" + str(dst)[:120]}
        return {"ok": True, "name": os.path.basename(dst)}

    def api_codex_auto_voice(self, name, lang="ja"):
        """素材供应链：一键配音——为剧本里所有有 text 的台词批量合成配音，
        写入包内 voice/ 并更新剧本引用。lang: ja(日文) / zh(中文，需声库支持)。
        返回 {ok, done, msg}。"""
        pkg = self._codex_pkg_path(name)
        jp = os.path.join(pkg, "codex.json")
        try:
            with open(jp, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            return {"ok": False, "err": "读剧本失败：" + str(e)[:120]}
        # UTAU 插件
        p = None
        try:
            if self.plugin_manager:
                p = self.plugin_manager.get_plugin("UTAU 语音")
        except Exception:
            p = None
        if not p or not getattr(p, "enabled", False):
            return {"ok": False, "err": "未启用 UTAU 语音插件（设置 → 插件 → UTAU 语音）"}
        if not hasattr(p, "_speak"):
            return {"ok": False, "err": "UTAU 插件缺少合成接口"}
        voice_dir = os.path.join(pkg, "voice")
        os.makedirs(voice_dir, exist_ok=True)
        done = 0
        failed = 0
        scenes = data.get("scenes") or []
        for si, sc in enumerate(scenes):
            for li, ln in enumerate(sc.get("lines") or []):
                if not isinstance(ln, dict):
                    continue
                text = str(ln.get("text") or "").strip()
                if not text or len(text) > 200:
                    continue
                try:
                    res = p._speak(text)
                    # _speak 播放并返回 "🗣️ text"；找到刚生成的 wav（tts_cache 最新）
                    base = getattr(app_paths, "get_base_dir", lambda: BASE_DIR)()
                    cache = os.path.join(base, "tts_cache")
                    cand = []
                    if os.path.isdir(cache):
                        cand = sorted([os.path.join(cache, f) for f in os.listdir(cache)
                                       if f.endswith(".wav")], key=os.path.getmtime, reverse=True)
                    if not cand:
                        failed += 1
                        continue
                    src = cand[0]
                    fn = f"auto_{si}_{li}.wav"
                    import shutil
                    shutil.copy2(src, os.path.join(voice_dir, fn))
                    ln["voice"] = "voice/" + fn
                    done += 1
                except Exception as e:
                    failed += 1
                    print(f"[CODEX voice] {si}/{li} 失败: {e}")
        save_guard.atomic_write_json(jp, data)
        return {"ok": True, "done": done, "failed": failed,
                "msg": f"配音完成：{done} 句成功" + (f"，{failed} 句失败" if failed else "")}

    def api_codex_draft(self, name, idea, count=3):
        """素材供应链：AI 起草剧本——给一句话主题，生成完整 codex.json。
        返回 {ok, script}（不落盘，前端预览后可保存）。"""
        idea = (idea or "").strip()
        if not idea:
            return {"ok": False, "err": "empty idea"}
        client = getattr(self.core, "client", None)
        if not client:
            return {"ok": False, "err": "未配置模型（设置里填 API Key）"}
        try:
            count = max(1, min(5, int(count or 3)))
        except (TypeError, ValueError):
            count = 3
        sys_prompt = (
            "你是视觉小说（GALGAME）剧本作家。根据用户给出的主题，生成一份 CODEX 剧本 JSON。\n"
            "格式（严格遵循，只输出 JSON）：\n"
            "{\n"
            '  "name": "剧本名",\n'
            '  "intro": "一句话简介",\n'
            '  "scenes": [\n'
            '    {"id": "s1", "title": "章节名", "bg": "", "lines": [\n'
            '      {"note": "舞台说明（旁白，无 speaker）"},\n'
            '      {"speaker": "角色名", "text": "台词"},\n'
            '      {"speaker": "角色名", "sprite": "", "text": "带动作的台词"},\n'
            '      {"choice": [{"text": "选项A", "goto": "s2"}, {"text": "选项B", "goto": "s3"}]},\n'
            '      {"end": "结局标题"}\n'
            '    ]}\n'
            '  ]\n'
            "}\n"
            "要求：\n"
            f"1. 生成 {count} 个场景（s1 起，id 连续），至少 2 个分支结局。\n"
            "2. 台词口语化、有情感张力，符合视觉小说节奏；旁白用 note 行。\n"
            "3. 选项分支用 goto 指向对应场景；场景末尾可 jump 回主线或用 end 结束。\n"
            "4. 只输出 JSON，不要 markdown 代码块包裹，不要额外说明。\n"
        )
        resp = client.chat.completions.create(
            model=self.core.model,
            messages=[{"role": "system", "content": sys_prompt},
                      {"role": "user", "content": idea}],
            stream=False, timeout=90,
        )
        raw = (resp.choices[0].message.content or "").strip()
        # 宽容解析：去代码块、取 JSON 对象
        import re as _re
        m = _re.search(r"\{[\s\S]*\}", raw)
        if not m:
            return {"ok": False, "err": "模型未输出有效 JSON"}
        try:
            data = json.loads(m.group(0))
        except Exception as e:
            return {"ok": False, "err": "JSON 解析失败：" + str(e)[:100]}
        if not isinstance(data, dict) or not data.get("scenes"):
            return {"ok": False, "err": "剧本缺少 scenes"}
        data.setdefault("name", (name or "我的故事"))
        data.setdefault("author", "")
        ok_v, issues = codex_core.validate_codex(data)
        return {"ok": True, "script": data, "valid": ok_v, "issues": issues}


    def api_codex_import_folder(self):
        """文件对话框选文件夹 → 自动归类为 CODEX 包（傻瓜化导入）"""
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            res = win.create_file_dialog(webview.FileDialog.FOLDER)
        except Exception as e:
            return {"ok": False, "err": "对话框失败：" + str(e)}
        if not res:
            return {"ok": False, "err": "cancelled"}
        src = res[0]
        base = os.path.basename(src.rstrip("/\\")) or "未命名"
        r = codex_core.import_folder(src, self.codex_dir, base)
        return {"ok": r["ok"], "msg": r.get("msg", ""), "name": base,
                "moved": r.get("moved", {})}

    def api_codex_get(self, name):
        """读取剧本 codex.json（JSON 功能：编辑预填）"""
        jp = os.path.join(self._codex_pkg_path(name), "codex.json")
        try:
            with open(jp, "r", encoding="utf-8") as f:
                data = json.load(f)
            ok, issues = codex_core.validate_codex(data)
            return {"ok": True, "script": data, "valid": ok, "issues": issues}
        except FileNotFoundError:
            return {"ok": False, "err": "not found"}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_codex_save(self, name, script_json):
        """保存剧本（校验 + 原子写盘）"""
        try:
            data = json.loads(script_json or "{}")
        except Exception as e:
            return {"ok": False, "err": "JSON 解析失败：" + str(e)[:120]}
        if not isinstance(data, dict):
            return {"ok": False, "err": "剧本必须是 JSON 对象"}
        data.setdefault("name", name)
        ok, issues = codex_core.validate_codex(data)
        jp = os.path.join(self._codex_pkg_path(name), "codex.json")
        try:
            save_guard.atomic_write_json(jp, data)
        except Exception as e:
            return {"ok": False, "err": "写入失败：" + str(e)[:120]}
        return {"ok": True, "valid": ok, "issues": issues}

    def api_codex_validate(self, name):
        """重新校验剧本，返回 (valid, issues)"""
        jp = os.path.join(self._codex_pkg_path(name), "codex.json")
        try:
            with open(jp, "r", encoding="utf-8") as f:
                data = json.load(f)
            ok, issues = codex_core.validate_codex(data)
            return {"ok": True, "valid": ok, "issues": issues}
        except Exception as e:
            return {"ok": False, "err": str(e)[:120]}

    def api_codex_delete(self, name):
        """删除 CODEX 包（整个目录）"""
        import shutil
        d = self._codex_pkg_path(name)
        if os.path.isdir(d):
            try:
                shutil.rmtree(d)
            except Exception as e:
                return {"ok": False, "err": str(e)[:120]}
        return {"ok": True}

    def api_codex_export(self, name):
        """导出 CODEX 包为 zip（JSON 功能：分享/备份）"""
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            safe = codex_core._safe_name(name)
            res = win.create_file_dialog(webview.FileDialog.SAVE,
                                         save_filename=safe + ".zip",
                                         file_types=("CODEX 包 (*.zip)",))
            if not res:
                return {"ok": False, "err": "cancelled"}
            path = res[0]
            if not path.lower().endswith(".zip"):
                path += ".zip"
            if codex_core.export_zip(self._codex_pkg_path(name), path):
                return {"ok": True, "file": os.path.basename(path)}
            return {"ok": False, "err": "导出失败"}
        except Exception as e:
            return {"ok": False, "err": "导出失败：" + str(e)[:120]}

    def api_codex_import_zip(self):
        """导入 CODEX zip 包"""
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            res = win.create_file_dialog(
                webview.FileDialog.OPEN, allow_multiple=False,
                file_types=("CODEX 包 (*.zip)", "All files (*.*)"))
        except Exception as e:
            return {"ok": False, "err": "对话框失败：" + str(e)}
        if not res:
            return {"ok": False, "err": "cancelled"}
        ok, name, msg = codex_core.import_zip(res[0], self.codex_dir)
        return {"ok": ok, "name": name, "msg": msg}

    def api_codex_pack(self, name):
        """打包：把 CODEX 包生成独立单文件 HTML（GALGAME，双击即玩，可分发）。
        默认保存到包内 dist/ 目录，返回可直接打开/分发的路径。"""
        pkg = self._codex_pkg_path(name)
        dist_dir = os.path.join(pkg, "dist")
        try:
            os.makedirs(dist_dir, exist_ok=True)
            safe = codex_core._safe_name(name)
            out = os.path.join(dist_dir, safe + ".html")
            ok, path, counts = codex_core.build_standalone_file(pkg, out)
            if not ok:
                return {"ok": False, "err": "打包失败"}
            size = os.path.getsize(path)
            return {"ok": True, "path": path, "size": size,
                    "counts": counts, "rel": os.path.relpath(path, pkg).replace("\\", "/")}
        except Exception as e:
            return {"ok": False, "err": "打包失败：" + str(e)[:200]}

    def api_codex_pack_save_as(self, name):
        """打包并另存为（另存对话框选位置）"""
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            safe = codex_core._safe_name(name)
            res = win.create_file_dialog(webview.FileDialog.SAVE,
                                         save_filename=safe + ".html",
                                         file_types=("GALGAME (*.html)",))
            if not res:
                return {"ok": False, "err": "cancelled"}
            out = res[0]
            if not out.lower().endswith(".html"):
                out += ".html"
            ok, path, counts = codex_core.build_standalone_file(self._codex_pkg_path(name), out)
            if not ok:
                return {"ok": False, "err": "打包失败"}
            return {"ok": True, "path": path,
                    "size": os.path.getsize(path), "counts": counts}
        except Exception as e:
            return {"ok": False, "err": "打包失败：" + str(e)[:200]}

    # ---------- 打包 EXE（独立可执行，双击即玩） ----------
    def api_codex_pack_exe(self, name):
        """打包 CODEX 包为独立 EXE（后台线程，完成后可 api_codex_pack_exe_status 查询）。
        需要本机 Python 环境含 PyInstaller。输出到 codex/<包名>/dist/。"""
        import threading
        try:
            pkg = self._codex_pkg_path(name)
            if not os.path.isfile(os.path.join(pkg, "codex.json")):
                return {"ok": False, "err": "包不存在"}
            state_file = os.path.join(pkg, "dist", ".exe_build_state.json")
            try:
                os.makedirs(os.path.dirname(state_file), exist_ok=True)
                with open(state_file, "w", encoding="utf-8") as f:
                    json.dump({"state": "building", "msg": "开始打包…"}, f)
            except Exception:
                pass
            threading.Thread(target=self._codex_exe_worker, args=(name,), daemon=True).start()
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def _codex_exe_worker(self, name):
        """后台：打包 EXE + 写状态文件"""
        state_file = os.path.join(self._codex_pkg_path(name), "dist", ".exe_build_state.json")
        def set_state(st, msg, extra=None):
            d = {"state": st, "msg": msg}
            if extra:
                d.update(extra)
            try:
                with open(state_file, "w", encoding="utf-8") as f:
                    json.dump(d, f, ensure_ascii=False)
            except Exception:
                pass
        try:
            set_state("building", "正在生成独立播放器…")
            pkg = self._codex_pkg_path(name)
            out_dir = os.path.join(pkg, "dist")
            # 找可用的 python 解释器（优先当前；冻结运行时退回 PATH 里的 python）
            import sys as _sys
            py_cmd = None
            if not getattr(_sys, "frozen", False):
                py_cmd = _sys.executable
            set_state("building", "正在调用 PyInstaller（首次约 1-2 分钟）…")
            exe, counts = codex_core.build_standalone_exe(
                pkg, out_dir, name=name, py_cmd=py_cmd,
                log_cb=lambda m: set_state("building", m))
            set_state("done", "打包完成",
                      {"exe": os.path.basename(exe), "size": os.path.getsize(exe),
                       "counts": counts,
                       "rel": os.path.relpath(exe, pkg).replace("\\", "/")})
        except Exception as e:
            set_state("error", str(e)[:300])

    def api_codex_pack_exe_status(self, name):
        """查询打包 EXE 进度（前端轮询）"""
        state_file = os.path.join(self._codex_pkg_path(name), "dist", ".exe_build_state.json")
        try:
            with open(state_file, "r", encoding="utf-8") as f:
                return {"ok": True, **json.load(f)}
        except FileNotFoundError:
            return {"ok": True, "state": "idle", "msg": "尚未开始"}
        except Exception as e:
            return {"ok": False, "err": str(e)[:120]}

    # ================= CODEX 系统权限（DICK 内置特权引擎） =================

    def api_codex_register_assoc(self):
        """注册 .codex 文件关联：双击 .codex 包 → 用 DICK 打开播放。
        写 HKCU 注册表（无需管理员）。返回 (ok, msg)。"""
        try:
            import winreg
            exe = os.path.abspath(sys.executable) if getattr(sys, "frozen", False) else \
                os.path.join(BASE_DIR, "DICK-HTML.exe")
            if not os.path.isfile(exe):
                exe = sys.executable
            # 1) ProgID
            with winreg.CreateKey(winreg.HKEY_CURRENT_USER,
                                  r"Software\Classes\DICK.Codex\shell\open\command") as k:
                winreg.SetValue(k, "", winreg.REG_SZ, f'"{exe}" "%1"')
            with winreg.CreateKey(winreg.HKEY_CURRENT_USER,
                                  r"Software\Classes\DICK.Codex\DefaultIcon") as k:
                winreg.SetValue(k, "", winreg.REG_SZ, f'"{exe}",0')
            # 2) 扩展名关联
            with winreg.CreateKey(winreg.HKEY_CURRENT_USER,
                                  r"Software\Classes\.codex") as k:
                winreg.SetValue(k, "", winreg.REG_SZ, "DICK.Codex")
            return {"ok": True, "msg": "✅ 已注册：双击 .codex 包将直接用 DICK 打开播放"}
        except Exception as e:
            return {"ok": False, "err": "注册失败：" + str(e)[:200]}

    def api_codex_export_codex(self, name):
        """导出 .codex 包（系统权限：保存到用户选的任意位置，双击即可在 DICK 打开）"""
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            safe = codex_core._safe_name(name)
            res = win.create_file_dialog(webview.FileDialog.SAVE,
                                         save_filename=safe + ".codex",
                                         file_types=("CODEX 包 (*.codex)",))
            if not res:
                return {"ok": False, "err": "cancelled"}
            path = res[0]
            if not path.lower().endswith(".codex"):
                path += ".codex"
            if codex_core.export_zip(self._codex_pkg_path(name), path):
                return {"ok": True, "file": os.path.basename(path)}
            return {"ok": False, "err": "导出失败"}
        except Exception as e:
            return {"ok": False, "err": "导出失败：" + str(e)[:200]}

    def api_codex_import_codex(self):
        """导入 .codex 包（系统权限：从任意位置选文件，自动解包到 codex/）"""
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            res = win.create_file_dialog(
                webview.FileDialog.OPEN, allow_multiple=False,
                file_types=("CODEX 包 (*.codex)", "All files (*.*)"))
        except Exception as e:
            return {"ok": False, "err": "对话框失败：" + str(e)}
        if not res:
            return {"ok": False, "err": "cancelled"}
        ok, name, msg = codex_core.import_zip(res[0], self.codex_dir)
        return {"ok": ok, "name": name, "msg": msg}

    def api_codex_run_program(self, cmd):
        """系统权限：执行本地程序/命令（action 钩子扩展）。
        cmd 是命令字符串，如 notepad.exe 或 完整路径。仅限一次调用，不做 shell 拼接注入。"""
        cmd = (cmd or "").strip()
        if not cmd:
            return {"ok": False, "err": "empty"}
        # 安全检查：只允许 exe/bat/cmd/vbs 直接执行；含 & | ; 等拼接符拒绝
        import re as _re
        if _re.search(r"[&|;><`]", cmd):
            return {"ok": False, "err": "禁止命令拼接"}
        low = cmd.lower().split()
        first = low[0] if low else ""
        if not (first.endswith(".exe") or first.endswith(".bat") or first.endswith(".cmd")
                or first in ("notepad", "calc", "mspaint", "explorer", "cmd", "start")):
            return {"ok": False, "err": "仅允许可执行程序"}
        try:
            import subprocess
            subprocess.Popen(cmd, shell=False)
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_codex_volume(self, level):
        """多媒体权限：设置播放器音量（0-100）。返回当前音量。"""
        try:
            level = max(0, min(100, int(level)))
        except (TypeError, ValueError):
            level = 100
        self.codex_volume = level
        return {"ok": True, "volume": level}

    def api_codex_fullscreen(self):
        """多媒体权限：CODEX 播放器全屏独占（pywebview 窗口切换）"""
        try:
            import webview
            win = webview.windows[0] if getattr(webview, "windows", None) else None
            if win is None:
                return {"ok": False, "err": "窗口未就绪"}
            win.toggle_fullscreen()
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def api_codex_clear_auto_open(self):
        """清除双击打开标记（前端播放器启动后调用，避免重复弹窗）"""
        self.codex_auto_open = None
        return {"ok": True}


    def api_codex_release(self, name):
        """两步发布（决战兵器·分发闭环）：一键打包 EXE + HTML + .codex 三件套到包内 dist/。
        后台线程执行，api_codex_pack_exe_status 轮询进度。"""
        import threading
        try:
            pkg = self._codex_pkg_path(name)
            if not os.path.isfile(os.path.join(pkg, "codex.json")):
                return {"ok": False, "err": "包不存在"}
            state_file = os.path.join(pkg, "dist", ".exe_build_state.json")
            try:
                os.makedirs(os.path.dirname(state_file), exist_ok=True)
                with open(state_file, "w", encoding="utf-8") as f:
                    json.dump({"state": "building", "msg": "开始发布…"}, f)
            except Exception:
                pass
            threading.Thread(target=self._codex_release_worker, args=(name,), daemon=True).start()
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)[:200]}

    def _codex_release_worker(self, name):
        """后台：发布三件套（HTML → EXE → .codex）"""
        import sys as _sys
        state_file = os.path.join(self._codex_pkg_path(name), "dist", ".exe_build_state.json")
        def set_state(st, msg, extra=None):
            d = {"state": st, "msg": msg}
            if extra:
                d.update(extra)
            try:
                with open(state_file, "w", encoding="utf-8") as f:
                    json.dump(d, f, ensure_ascii=False)
            except Exception:
                pass
        try:
            pkg = self._codex_pkg_path(name)
            out_dir = os.path.join(pkg, "dist")
            os.makedirs(out_dir, exist_ok=True)
            safe = codex_core._safe_name(name)
            set_state("building", "① 生成独立 HTML…")
            ok, path, counts = codex_core.build_standalone_file(pkg, os.path.join(out_dir, safe + ".html"))
            if not ok:
                set_state("error", "HTML 打包失败")
                return
            # ② EXE
            set_state("building", "② 打包 EXE（首次约 1-2 分钟）…")
            py_cmd = None if getattr(_sys, "frozen", False) else _sys.executable
            exe, counts = codex_core.build_standalone_exe(pkg, out_dir, name=name, py_cmd=py_cmd)
            # ③ .codex
            set_state("building", "③ 生成 .codex 包…")
            codex_path = os.path.join(out_dir, safe + ".codex")
            codex_core.export_zip(pkg, codex_path)
            set_state("done", "发布完成：EXE + HTML + .codex",
                      {"exe": os.path.basename(exe),
                       "html": safe + ".html",
                       "codex": safe + ".codex",
                       "size_exe": os.path.getsize(exe),
                       "size_html": os.path.getsize(os.path.join(out_dir, safe + ".html")),
                       "size_codex": os.path.getsize(codex_path),
                       "rel": "dist",
                       "counts": counts})
        except Exception as e:
            set_state("error", str(e)[:300])



    def api_codex_asset(self, name, kind, file):
        """返回 CODEX 包内资源的 dataURL（播放器/预览用）。
        kind: sprites/bg/bgm/voice；file: 相对该 kind 目录的文件名"""
        kind = (kind or "").lower()
        if kind not in codex_core.SUBDIRS:
            return {"ok": False, "err": "bad kind"}
        base = os.path.basename((file or "").replace("\\", "/"))
        if not base:
            return {"ok": False, "err": "bad file"}
        p = os.path.join(self._codex_pkg_path(name), kind, base)
        try:
            with open(p, "rb") as f:
                raw = f.read()
        except Exception:
            return {"ok": False, "err": "not found"}
        import base64 as _b64
        ext = os.path.splitext(base)[1].lower()
        mime = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                ".webp": "image/webp", ".gif": "image/gif",
                ".mp3": "audio/mpeg", ".wav": "audio/wav", ".ogg": "audio/ogg",
                ".flac": "audio/flac", ".mp4": "audio/mp4", ".m4a": "audio/mp4"}.get(ext, "application/octet-stream")
        return {"ok": True, "data": "data:" + mime + ";base64," + _b64.b64encode(raw).decode("ascii")}

    def api_codex_script_files(self, name):
        """列出包内全部资源文件（按 kind 分组）——播放器/编辑器引用列表"""
        out = {k: [] for k in codex_core.SUBDIRS}
        for kind in codex_core.SUBDIRS:
            kd = os.path.join(self._codex_pkg_path(name), kind)
            if os.path.isdir(kd):
                try:
                    out[kind] = sorted(x for x in os.listdir(kd)
                                       if os.path.isfile(os.path.join(kd, x)))
                except Exception:
                    pass
        return {"ok": True, "files": out}

    def api_speak_node(self, node_id):
        """消息旁喇叭按钮：按 node_id 朗读该条 AI 消息（合成缓存于 tts_cache，重复点击不重合成）。
        取 [ja] 日配句优先，无则用正文；走 UTAU/HANASU 双引擎。"""
        try:
            node = self.core.tree.nodes.get(node_id)
            if not node:
                return {"ok": False, "err": "节点不存在"}
            if node.role != "assistant":
                return {"ok": False, "err": "仅 AI 消息可朗读"}
            p = None
            try:
                if self.plugin_manager:
                    p = self.plugin_manager.get_plugin("UTAU 语音")
            except Exception:
                p = None
            if not p or not getattr(p, "enabled", False):
                return {"ok": False, "err": "未启用语音插件（设置 → 插件 → UTAU 语音）"}
            if not hasattr(p, "_speak"):
                return {"ok": False, "err": "语音插件缺少合成接口"}
            # 优先 [ja] 日配句
            text = ""
            try:
                ja = (node.metadata or {}).get("ja")
                if isinstance(ja, str) and ja.strip():
                    text = ja.strip()
            except Exception:
                pass
            if not text:
                text = str(node.content or "").strip()
            if not text:
                return {"ok": False, "err": "无内容可朗读"}
            # 合成缓存：按 node_id 复用
            base = getattr(app_paths, "get_base_dir", lambda: BASE_DIR)()
            cache = os.path.join(base, "tts_cache")
            os.makedirs(cache, exist_ok=True)
            cached = os.path.join(cache, "node_" + str(node_id) + ".wav")
            if os.path.exists(cached):
                p._play(cached)
                return {"ok": True, "cached": True}
            # 合成前预处理：中文 → 日文（走日文补丁翻译，日文声库才能念）
            final_text = text[:200]
            if hasattr(p, "_prepare_text"):
                final_text, _tr = p._prepare_text(final_text, cache)
            out, _eng = p._engine().synthesize(final_text, cache, pitch_mode="auto") if hasattr(p, "_engine") else (None, None)
            if out and os.path.exists(out):
                try:
                    import shutil
                    shutil.copyfile(out, cached)
                except Exception:
                    pass
                p._play(out)
                return {"ok": True}
            # 兼容旧插件（无 _engine 的 _speak 直接播放）
            res = p._speak(text[:200])
            if isinstance(res, str) and res.startswith("⚠️"):
                return {"ok": False, "err": res}
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)[:120]}

    def api_codex_run_action(self, cmd):
        """播放器行动钩子（CODEX 深度集成 DICK 的系统权限）。
        支持：
          plugin:<命令>   调用插件命令（如 plugin:/speak こんにちは）
          model:<文本>    调用模型生成一句话（如 model:请描述这个场景的天气）
          aff:<±N>        修改当前角色好感度（机制卡）
          run:<程序>      调用本地程序（如 run:notepad.exe）
          <其他>          默认按插件命令处理（兼容旧剧本 /speak 等）
        返回 {ok, out}。"""
        cmd = (cmd or "").strip()
        if not cmd:
            return {"ok": True, "out": ""}
        low = cmd.lower()
        try:
            if low.startswith("model:"):
                prompt = cmd[len("model:"):].strip()
                if not prompt:
                    return {"ok": True, "out": ""}
                client = getattr(self.core, "client", None)
                if not client:
                    return {"ok": True, "out": "（未配置模型）"}
                try:
                    resp = client.chat.completions.create(
                        model=self.core.model,
                        messages=[{"role": "user", "content": prompt}],
                        stream=False, timeout=30)
                    text = (resp.choices[0].message.content or "").strip()
                    return {"ok": True, "out": text[:200]}
                except Exception as e:
                    return {"ok": True, "out": "（模型调用失败：" + str(e)[:60] + "）"}
            if low.startswith("aff:"):
                try:
                    delta = int(cmd[len("aff:"):].strip())
                    st = getattr(self.core, "mechanism_state", None)
                    if st is None or "affection" not in st:
                        return {"ok": True, "out": "（未启用好感度）"}
                    cfg = getattr(self.core, "_mech_config", None) or {}
                    aff = cfg.get("affection") if isinstance(cfg.get("affection"), dict) else {}
                    hi = int(aff.get("max", 100) or 100)
                    lo = int(aff.get("min", 0) or 0)
                    cur = int(st.get("affection", 50) or 50)
                    st["affection"] = max(lo, min(hi, cur + delta))
                    return {"ok": True, "out": f"❤️ 好感 {cur}→{st['affection']}"}
                except Exception as e:
                    return {"ok": True, "out": "（好感调整失败）"}
            if low.startswith("run:"):
                prog = cmd[len("run:"):].strip()
                r = self.api_codex_run_program(prog)
                return {"ok": r.get("ok", False), "out": r.get("err", "已调用本地程序")}
            # 默认：插件命令（兼容 plugin: 前缀与旧裸命令）
            plugin_cmd = cmd[len("plugin:"):].strip() if low.startswith("plugin:") else cmd
            if self.plugin_manager:
                out = self.plugin_manager.handle_command(plugin_cmd)
                if isinstance(out, str):
                    return {"ok": True, "out": out}
                if isinstance(out, tuple) and out:
                    return {"ok": True, "out": str(out[0])}
        except Exception as e:
            print(f"[CODEX action] 执行失败: {e}")
        return {"ok": True, "out": ""}

    def _maybe_auto_turn(self, rounds):
        """群聊自动接话：当前回复后由其他角色接力发言（最多 rounds 轮）"""
        if rounds <= 0 or not self.auto_turn or len(self.selected_roles) < 2 or self.busy:
            return
        roster = [r["name"] for r in self.roles if r["name"] in self.selected_roles]
        last = getattr(self.core, "last_speaker", None)
        candidates = [n for n in roster if n != last]
        if not candidates:
            candidates = roster[1:] or roster
        speaker = candidates[0]
        self.busy = True
        self.streaming = ""

        def go():
            try:
                self.core.send_auto_turn(
                    speaker,
                    on_response=lambda reply, usage: self._auto_turn_done(reply, usage, rounds),
                    on_error=lambda err: self._auto_turn_fail())
            except Exception:
                self._auto_turn_fail()

        # 延迟启动：等上一轮 fetch 的 is_processing 复位，避免被“正在处理中”拦截
        threading.Timer(0.5, go).start()

    def _auto_turn_done(self, reply, usage, rounds):
        if usage:
            self.total_tokens += int(getattr(usage, "total_tokens", 0) or 0)
        self._save_tree()
        self._rebuild_messages()
        self.busy = False
        self._maybe_auto_turn(rounds - 1)

    def _auto_turn_fail(self):
        self.busy = False
        self.streaming = ""

    def api_toggle_language(self):
        self.language = "en" if self.language == "zh" else "zh"
        i18n.set_lang(self.language)
        self.config["language"] = self.language
        self._save_config()
        return {"ok": True, "lang": self.language}

    def api_set_theme(self, theme_idx, accent_idx):
        self.config["ui_theme"] = int(theme_idx)
        self.config["ui_accent"] = int(accent_idx)
        self._save_config()
        return {"ok": True}

    def api_save_key(self, key, base_url, model, provider_id=None):
        pid = (provider_id or self.provider_id or "deepseek").strip()
        self.provider_id = pid
        k = (key or "").strip()
        self.api_keys[pid] = k
        self.config["api_keys"] = self.api_keys
        self.config["provider"] = pid
        self.config["api_key"] = k  # 兼容旧字段
        self.config["base_url"] = (base_url or "").strip()
        self.config["model"] = (model or "").strip()
        # 无条件重建客户端：切换厂商时不能沿用旧厂商的 Key 与 URL
        # 免费厂商（OVH 免费链/Ollama）无 Key：用占位符构造客户端（其服务端忽略认证头）
        self.core.set_api_key(k if k else "free")
        if self.config["base_url"]:
            self.core.set_base_url(self.config["base_url"])
        if self.config["model"]:
            self.core.set_model(self.config["model"])
        # 重建客户端会丢失代理，重新应用
        if self.config.get("proxy"):
            self.core.set_proxy(self.config["proxy"])
        # 换厂商后停止序列按新模型家族默认生效
        self.core.set_stop_sequences(self._effective_stop())
        self._save_config()
        return {"ok": True, "provider": pid}

    def api_set_proxy(self, proxy):
        """设置 LLM 通道代理（http/https/socks5；空串 = 直连）"""
        proxy = (proxy or "").strip()
        self.config["proxy"] = proxy
        self.core.set_proxy(proxy)
        self._save_config()
        return {"ok": True}

    def api_set_relay(self, url):
        """设置内置代理通道地址；空串 = 关闭（恢复纯直连）"""
        url = ((url or "").strip().rstrip("/")) or ""
        self.config["relay_url"] = url
        self.core.relay_on = False          # 重置，回到「直连优先」模式
        self.core.set_relay(url or BUILTIN_RELAY)
        self._save_config()
        return {"ok": True}

    def api_open_url(self, url):
        """在系统浏览器打开模型商官网（注册/充值跳转）"""
        url = (url or "").strip()
        if not url.startswith("http://") and not url.startswith("https://"):
            return {"ok": False, "err": "invalid url"}
        try:
            import webbrowser
            webbrowser.open(url)
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "err": str(e)}

    def _save_config(self):
        try:
            save_guard.atomic_write_json(self.config_file, self.config)
        except Exception:
            pass


# ---------- 免费视觉链（传图补丁） ----------
_VISION_MODELS = ["Qwen3.5-397B-A17B", "Qwen2.5-VL-72B-Instruct", "Qwen3.6-27B",
                  "Mistral-Small-3.2-24B-Instruct-2506", "Qwen3.5-9B"]


def _vision_describe(image_b64, mime, proxies=None):
    import requests
    data_url = "data:" + mime + ";base64," + image_b64
    for model in _VISION_MODELS:
        body = {
            "model": model,
            "messages": [{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": data_url}},
                {"type": "text", "text": "请用中文详细描述这张图片的内容（包括文字、物体、场景、数据，如有表格请逐项列出）。"},
            ]}],
            "max_tokens": 4096,
            "stream": False,
        }
        try:
            r = requests.post("https://oai.endpoints.kepler.ai.cloud.ovh.net/v1/chat/completions",
                              json=body, headers={"Content-Type": "application/json"},
                              timeout=90, proxies=proxies)
            if r.status_code == 429:
                continue
            if r.status_code >= 400:
                continue
            content = ((r.json().get("choices") or [{}])[0].get("message") or {}).get("content")
            if isinstance(content, str) and content.strip():
                return content.strip()
        except Exception:
            continue
    return None


def main():
    import webview
    app = HtmlApp()
    # 系统权限：双击 .codex 文件启动 → 自动导入并打开 CODEX 播放器
    if len(sys.argv) > 1:
        p = sys.argv[1]
        if p.lower().endswith(".codex") and os.path.isfile(p):
            try:
                ok, name, msg = codex_core.import_zip(p, app.codex_dir)
                if ok:
                    app.codex_auto_open = name
                    print(f"[CODEX] 双击打开: {name}（{msg}）")
            except Exception as e:
                print(f"[CODEX] 打开失败: {e}")
    html_path = os.path.join(_web_root(), "index.html")
    window = webview.create_window(
        "Direct-Interface Cork-bore Kit v2.0", html_path, js_api=app,
        width=1060, height=820, min_size=(860, 640), background_color="#0f1115")
    webview.start()


if __name__ == "__main__":
    main()
